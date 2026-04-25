#!/usr/bin/env python3.11
"""
astro-natal-merkaba: render_solar_docx.py
DOCX-отчёт по соляру (палитра МерКаБа).
"""
import argparse
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("❌ python-docx не установлен", file=sys.stderr)
    sys.exit(1)

from render_docx import (
    setup_styles, add_hr, add_page_break, set_cell_width,
    COLOR_TITLE, COLOR_H1, COLOR_H2, COLOR_MUTED,
)
from render_transits_docx import (
    INTENSITY_COLOR, NATURE_LABEL, format_date_ru,
)

HOUSE_THEMES = {
    1: 'Я, тело, самопроявление',
    2: 'Деньги, ресурсы, ценности',
    3: 'Коммуникация, окружение, обучение',
    4: 'Дом, семья, корни',
    5: 'Творчество, дети, удовольствия',
    6: 'Здоровье, работа, рутина',
    7: 'Партнёрство, отношения',
    8: 'Трансформация, общие ресурсы',
    9: 'Мировоззрение, путешествия, философия',
    10: 'Карьера, общественная роль',
    11: 'Друзья, цели, будущее',
    12: 'Уединение, психея, скрытое',
}


def add_header(doc, solar):
    meta = solar['meta']
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f"Соляр {meta['natal_name']} — {meta['solar_year']}")
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = COLOR_TITLE
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"Точка возврата: {meta.get('solar_moment_human', '?')}")
    sr.font.size = Pt(14)
    sr.font.color.rgb = COLOR_H2
    sub.paragraph_format.space_after = Pt(8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(
        f"Натал: {format_date_ru(meta.get('natal_date'))} · "
        f"Город соляра: {meta.get('solar_city', '?')}"
    )
    ir.font.size = Pt(10)
    ir.font.color.rgb = COLOR_MUTED
    ir.italic = True
    info.paragraph_format.space_after = Pt(12)

    add_hr(doc)


def add_how_to_read(doc, text):
    if not text:
        return
    doc.add_heading('Как читать этот отчёт', level=1)
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_year_summary(doc, solar, summary_text):
    doc.add_heading('Тема года', level=1)
    sun_house = solar.get('sun_in_house')
    sun_pos = solar.get('solar_planets', {}).get('sun', {})
    asc = solar.get('solar_ascendant', {})
    mc = solar.get('solar_mc', {})

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ('☉ Солнце соляра',
         f"{sun_pos.get('sign_ru', '?')} {sun_pos.get('degrees', '?')}° · "
         f"дом {sun_house} ({HOUSE_THEMES.get(sun_house, '—')})"),
        ('↑ Асцендент соляра',
         f"{asc.get('sign_ru', '?')} {asc.get('degrees', '?')}°"),
        ('⊕ MC соляра',
         f"{mc.get('sign_ru', '?')} {mc.get('degrees', '?')}°"),
        ('Угловых планет', str(solar.get('totals', {}).get('angular_count', 0))),
    ]
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        for r in cells[0].paragraphs[0].runs:
            r.font.bold = True
        set_cell_width(cells[0], 2.0)
        set_cell_width(cells[1], 4.5)

    if summary_text:
        p = doc.add_paragraph(summary_text)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(10)


def add_angular_planets(doc, solar, angular_text):
    angular = solar.get('angular_planets', [])
    if not angular:
        return
    add_page_break(doc)
    doc.add_heading('Угловые планеты — ключевые силы года', level=1)
    p = doc.add_paragraph(
        'Планеты в 1, 4, 7 и 10 домах соляра — это самые активные силы года. '
        'Их архетипы будут особенно ощутимы и проявятся в реальной жизни сильнее всего.'
    )
    p.paragraph_format.space_after = Pt(10)

    if angular_text:
        p = doc.add_paragraph(angular_text)
        p.paragraph_format.space_after = Pt(10)

    for ap in angular:
        h2 = doc.add_paragraph()
        h2.style = doc.styles['Heading 2']
        h2.paragraph_format.keep_with_next = True
        title_run = h2.add_run(
            f"{ap['symbol']} {ap['name_ru']} в доме {ap['house']} "
            f"({HOUSE_THEMES.get(ap['house'], '—')})"
        )
        title_run.font.color.rgb = COLOR_H2

        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.keep_with_next = True
        text = f"{ap['sign_ru']} {ap['degrees']}°"
        if ap.get('retrograde'):
            text += ' ℞'
        mr = meta_p.add_run(text)
        mr.font.size = Pt(10)
        mr.italic = True
        mr.font.color.rgb = COLOR_MUTED


def add_house_distribution(doc, solar, house_text):
    """Какие натальные планеты в каких домах соляра."""
    planets = solar.get('solar_planets', {})
    doc.add_heading('Планеты по домам соляра', level=1)
    if house_text:
        p = doc.add_paragraph(house_text)
        p.paragraph_format.space_after = Pt(10)

    by_house = {}
    for k, p_data in planets.items():
        h = p_data.get('house')
        if h is None:
            continue
        by_house.setdefault(h, []).append(p_data)

    for house_num in range(1, 13):
        plist = by_house.get(house_num, [])
        if not plist:
            continue
        h2 = doc.add_paragraph()
        h2.style = doc.styles['Heading 2']
        h2.paragraph_format.keep_with_next = True
        h2.add_run(f"Дом {house_num} — {HOUSE_THEMES.get(house_num, '?')}")

        line = ', '.join(
            f"{p['symbol']} {p['name_ru']} ({p['sign_ru']} {p['degrees']}°"
            + (' ℞' if p.get('retrograde') else '') + ')'
            for p in plist
        )
        para = doc.add_paragraph(line)
        para.paragraph_format.space_after = Pt(6)


def add_aspects_table(doc, solar, aspects_text):
    aspects = solar.get('aspects', [])
    if not aspects:
        return
    add_page_break(doc)
    doc.add_heading('Аспекты внутри соляра', level=1)
    if aspects_text:
        p = doc.add_paragraph(aspects_text)
        p.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=len(aspects) + 1, cols=3)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [3.0, 1.5, 1.5]
    for i, h in enumerate(['Аспект', 'Природа', 'Орб']):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_width(cell, widths[i])
        for r in cell.paragraphs[0].runs:
            r.font.bold = True

    for row_i, asp in enumerate(aspects, start=1):
        cells = table.rows[row_i].cells
        cells[0].text = (
            f"{asp['planet1_symbol']} {asp['planet1_ru']} "
            f"{asp['aspect_symbol']} {asp['planet2_symbol']} {asp['planet2_ru']}"
        )
        cells[1].text = NATURE_LABEL.get(asp.get('aspect_nature', 'neutral'), '—')
        cells[2].text = f"{asp['orb']}°"
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_practical_advice(doc, advice_list):
    if not advice_list:
        return
    add_page_break(doc)
    doc.add_heading('Что важно в этот год', level=1)
    p = doc.add_paragraph(
        'Соляр описывает энергетический ландшафт года. Эти ориентиры — '
        'про то, на что направить внимание и где раскрыть потенциал.'
    )
    p.paragraph_format.space_after = Pt(8)
    for advice in advice_list:
        bullet = doc.add_paragraph(advice, style='List Bullet')
        bullet.paragraph_format.space_after = Pt(4)


def add_colophon(doc, solar):
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'Соляр рассчитан через Swiss Ephemeris (kerykeion). Точка возврата '
        'найдена бинарным поиском с точностью ~1 минута. Город пребывания в день ДР '
        'этого года определяет дома и углы соляра — переезд между двумя ДР меняет «дом» соляра.'
    )
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED
    r.italic = True

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run('astro-natal-merkaba · автор: Дмитрий · dimkaklasnyi@gmail.com')
    ar.font.size = Pt(8)
    ar.font.color.rgb = COLOR_MUTED


def render(solar_path, interp_path, out_path):
    with open(solar_path, 'r', encoding='utf-8') as f:
        solar = json.load(f)
    if interp_path and Path(interp_path).exists():
        with open(interp_path, 'r', encoding='utf-8') as f:
            interp = json.load(f)
    else:
        interp = {}

    doc = Document()
    setup_styles(doc)

    add_header(doc, solar)
    add_how_to_read(doc, interp.get('intro_how_to_read'))
    add_year_summary(doc, solar, interp.get('summary'))
    add_angular_planets(doc, solar, interp.get('angular'))
    add_house_distribution(doc, solar, interp.get('houses'))
    add_aspects_table(doc, solar, interp.get('aspects'))
    add_practical_advice(doc, interp.get('practical_advice', []))
    add_colophon(doc, solar)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main():
    p = argparse.ArgumentParser()
    p.add_argument('--solar', required=True)
    p.add_argument('--interp')
    p.add_argument('--out', required=True)
    args = p.parse_args()
    out = render(args.solar, args.interp, args.out)
    print(f"  📄 DOCX: {out}")


if __name__ == '__main__':
    main()
