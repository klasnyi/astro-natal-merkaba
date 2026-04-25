#!/usr/bin/env python3.11
"""
astro-natal-merkaba: render_synastry_docx.py
DOCX-отчёт по синастрии (палитра МерКаБа).
"""
import argparse
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("❌ python-docx не установлен", file=sys.stderr)
    sys.exit(1)

from render_docx import (
    setup_styles, add_hr, add_page_break, set_cell_width,
    COLOR_TITLE, COLOR_H1, COLOR_H2, COLOR_MUTED,
)
from render_transits_docx import (
    INTENSITY_COLOR, NATURE_LABEL, format_date_ru,
)
from render_solar_docx import HOUSE_THEMES


GREEN = RGBColor(0x27, 0xAE, 0x60)
RED = RGBColor(0xC0, 0x39, 0x2B)


def add_header(doc, syn):
    meta = syn['meta']
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f"Синастрия: {meta['person1_name']} × {meta['person2_name']}")
    tr.font.size = Pt(24)
    tr.font.bold = True
    tr.font.color.rgb = COLOR_TITLE
    title.paragraph_format.space_after = Pt(6)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(
        f"{meta['person1_name']}: {format_date_ru(meta.get('person1_date'))} · "
        f"{meta.get('person1_city', '?')}\n"
        f"{meta['person2_name']}: {format_date_ru(meta.get('person2_date'))} · "
        f"{meta.get('person2_city', '?')}"
    )
    ir.font.size = Pt(10)
    ir.font.color.rgb = COLOR_MUTED
    ir.italic = True
    info.paragraph_format.space_after = Pt(12)
    add_hr(doc)


def add_how_to_read(doc, text):
    if not text:
        return
    doc.add_heading('Как читать этот отчёт', level=1)
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_overview_summary(doc, syn, summary_text):
    doc.add_heading('Общая динамика отношений', level=1)
    if summary_text:
        p = doc.add_paragraph(summary_text)
        p.paragraph_format.space_after = Pt(10)

    totals = syn.get('totals', {})
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ('Всего cross-аспектов', str(totals.get('cross_aspects_total', 0))),
        ('🟢 Гармоничных', str(totals.get('harmonious_count', 0))),
        ('🔴 Напряжённых', str(totals.get('tense_count', 0))),
        ('💫 Романтических', str(totals.get('romantic_aspects_count', 0))),
    ]
    for i, (lbl, val) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = lbl
        cells[1].text = val
        for r in cells[0].paragraphs[0].runs:
            r.font.bold = True
        set_cell_width(cells[0], 3.0)
        set_cell_width(cells[1], 1.5)


def add_top_aspects(doc, syn, aspects_interp):
    aspects = syn.get('cross_aspects_top', [])[:12]
    if not aspects:
        return
    add_page_break(doc)
    doc.add_heading('Ключевые контакты — детально', level=1)
    p = doc.add_paragraph(
        'Топ контактов между картами в порядке силы. Природа аспекта — '
        'это качество энергии: гармоничные дают поддержку, напряжённые — рост через трение, '
        'соединения — слияние. Орб ниже 1° — точечный контакт, очень ощутимый.'
    )
    p.paragraph_format.space_after = Pt(10)

    for asp in aspects:
        h2 = doc.add_paragraph()
        h2.style = doc.styles['Heading 2']
        h2.paragraph_format.keep_with_next = True
        title_run = h2.add_run(
            f"{asp['p1_symbol']} {asp['p1_ru']} ({asp['p1_owner']}) "
            f"{asp['aspect_symbol']} "
            f"{asp['p2_symbol']} {asp['p2_ru']} ({asp['p2_owner']})"
        )
        # Цвет по природе
        nature = asp.get('aspect_nature', 'neutral')
        if nature == 'tense':
            title_run.font.color.rgb = RED
        elif nature == 'harmonious':
            title_run.font.color.rgb = GREEN
        else:
            title_run.font.color.rgb = COLOR_H2

        # Подзаголовок
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.keep_with_next = True
        text = (
            f"{asp['p1_ru']} ({asp['p1_owner']}) в {asp['p1_sign_ru']} {asp['p1_degrees']}° "
            f"⇄ {asp['p2_ru']} ({asp['p2_owner']}) в {asp['p2_sign_ru']} {asp['p2_degrees']}°"
        )
        text += f" · орб {asp['orb']}° · {NATURE_LABEL.get(nature, '—')}"
        if asp.get('is_romantic'):
            text += ' · 💫 романтический индикатор'
        mr = meta_p.add_run(text)
        mr.font.size = Pt(10)
        mr.italic = True
        mr.font.color.rgb = COLOR_MUTED
        meta_p.paragraph_format.space_after = Pt(4)

        text = aspects_interp.get(asp['key'])
        if text:
            tp = doc.add_paragraph(text)
            tp.paragraph_format.space_after = Pt(10)
        else:
            tp = doc.add_paragraph(
                f"(Интерпретация для {asp['p1_ru']} {asp['aspect_ru']} {asp['p2_ru']} "
                f"не сгенерирована — заполни поле '{asp['key']}' в interp.json.)"
            )
            tp.runs[0].italic = True
            tp.runs[0].font.color.rgb = COLOR_MUTED


def add_house_overlay(doc, overlays, owner_from, owner_into, overlay_text):
    if not overlays:
        return
    doc.add_heading(f'Где планеты {owner_from} в домах {owner_into}', level=1)
    if overlay_text:
        p = doc.add_paragraph(overlay_text)
        p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph(
        f"Когда планеты одного человека «падают» в дома другого — они активизируют "
        f"конкретные сферы жизни этого человека. Например, планета в 7-м доме партнёра "
        f"касается темы партнёрства; в 10-м — карьеры и общественной роли."
    )
    p.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=len(overlays) + 1, cols=3)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [2.2, 2.0, 2.5]
    for i, h in enumerate(['Планета', 'В доме', 'Тема дома']):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_width(cell, widths[i])
        for r in cell.paragraphs[0].runs:
            r.font.bold = True

    for row_i, ov in enumerate(overlays, start=1):
        cells = table.rows[row_i].cells
        cells[0].text = f"{ov['symbol']} {ov['planet_ru']} ({ov['sign_ru']} {ov['degrees']}°)"
        cells[1].text = f"Дом {ov['falls_in_house']}"
        cells[2].text = HOUSE_THEMES.get(ov['falls_in_house'], '—')
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_practical_advice(doc, advice_list):
    if not advice_list:
        return
    add_page_break(doc)
    doc.add_heading('Что важно понять про эти отношения', level=1)
    p = doc.add_paragraph(
        'Синастрия описывает энергетический потенциал отношений — не приговор. '
        'Любые контакты можно прожить осознанно через диалог и работу с собой.'
    )
    p.paragraph_format.space_after = Pt(8)
    for advice in advice_list:
        bullet = doc.add_paragraph(advice, style='List Bullet')
        bullet.paragraph_format.space_after = Pt(4)


def add_colophon(doc, syn):
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'Синастрия рассчитана через Swiss Ephemeris (kerykeion). Орбы синастрии: '
        '5-6° для личных планет, 4° для социальных и трансперсональных. Природа аспекта '
        '— классическая (соединение нейтрально, секстиль/трин гармоничны, квадрат/оппозиция напряжены). '
        'Рекомендации архетипические, не предсказания событий.'
    )
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED
    r.italic = True

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run('astro-natal-merkaba · автор: Дмитрий · dimkaklasnyi@gmail.com')
    ar.font.size = Pt(8)
    ar.font.color.rgb = COLOR_MUTED


def render(syn_path, interp_path, out_path):
    with open(syn_path, 'r', encoding='utf-8') as f:
        syn = json.load(f)
    if interp_path and Path(interp_path).exists():
        with open(interp_path, 'r', encoding='utf-8') as f:
            interp = json.load(f)
    else:
        interp = {}

    doc = Document()
    setup_styles(doc)

    add_header(doc, syn)
    add_how_to_read(doc, interp.get('intro_how_to_read'))
    add_overview_summary(doc, syn, interp.get('summary'))
    add_top_aspects(doc, syn, interp.get('aspects', {}))

    name1 = syn['meta']['person1_name']
    name2 = syn['meta']['person2_name']
    add_house_overlay(doc, syn.get('house_overlay_1_to_2', []),
                      name1, name2, interp.get('overlay_1_to_2'))
    add_house_overlay(doc, syn.get('house_overlay_2_to_1', []),
                      name2, name1, interp.get('overlay_2_to_1'))
    add_practical_advice(doc, interp.get('practical_advice', []))
    add_colophon(doc, syn)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main():
    p = argparse.ArgumentParser()
    p.add_argument('--synastry', required=True)
    p.add_argument('--interp')
    p.add_argument('--out', required=True)
    args = p.parse_args()
    out = render(args.synastry, args.interp, args.out)
    print(f"  📄 DOCX: {out}")


if __name__ == '__main__':
    main()
