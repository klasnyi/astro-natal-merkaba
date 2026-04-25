#!/usr/bin/env python3.11
"""
astro-natal-merkaba: render_docx.py
Рендерит стильный DOCX из натальной карты + LLM-интерпретаций.

Вход:
  --chart    <path>  chart JSON от build_chart.py
  --interp   <path>  interpretations JSON от LLM
  --out      <path>  путь DOCX (default: рядом с chart JSON)

Структура interpretations JSON:
{
  "intro_how_to_read": "...",
  "sun_moon_asc_summary": "...",
  "planets": {"sun": "...", "moon": "...", ...},
  "houses": {"1": "...", "2": "...", ...},        # опционально если has_time
  "key_aspects": {"sun_conjunction_mercury": "...", ...},
  "elements_modalities": "...",
  "stelliums": "..." | null,
  "vedic": {"moon_nakshatra": "...", "ascendant_nakshatra": "..."} | null,
  "practical_advice": ["совет 1", "совет 2", ...]
}
"""
import argparse
import json
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ python-docx не установлен: pip install python-docx")
    sys.exit(1)


# ─── Расширенные интерпретации (Phase 16, v2.4) ──────────────────────────────

_EXT_REFS_CACHE = None

def _load_extended_refs():
    """Lazy-load расширенных reference JSON. Возвращает dict с тремя ключами."""
    global _EXT_REFS_CACHE
    if _EXT_REFS_CACHE is not None:
        return _EXT_REFS_CACHE
    from pathlib import Path as _Path
    refs_dir = _Path(__file__).parent.parent / 'references'
    out = {}
    for name in ('planets_in_signs', 'planets_in_houses',
                 'elements_modalities_hemispheres'):
        path = refs_dir / f'{name}.json'
        try:
            if path.exists():
                with open(path, 'r', encoding='utf-8') as f:
                    out[name] = json.load(f)
            else:
                out[name] = {}
        except Exception:
            out[name] = {}
    _EXT_REFS_CACHE = out
    return out


def _format_ext_entry(entry):
    """{keyword, archetype, gift, shadow|focus} → одна строка."""
    if not entry or not isinstance(entry, dict):
        return None
    parts = []
    kw = entry.get('keyword')
    arch = entry.get('archetype')
    gift = entry.get('gift')
    shadow = entry.get('shadow')
    focus = entry.get('focus')
    if kw and arch:
        parts.append(f"«{kw}». {arch}")
    elif arch:
        parts.append(arch)
    elif kw:
        parts.append(kw)
    if focus:
        parts.append(f"Фокус: {focus}")
    if gift:
        parts.append(f"Дар: {gift}")
    if shadow:
        parts.append(f"Тень: {shadow}")
    return '  '.join(parts) if parts else None


def get_planet_in_sign_text(planet_key, sign_ru):
    """Возвращает расширенный текст планеты в знаке или None если нет."""
    refs = _load_extended_refs()
    entry = refs.get('planets_in_signs', {}).get(planet_key, {}).get(sign_ru)
    return _format_ext_entry(entry)


def get_planet_in_house_text(planet_key, house_num):
    """Возвращает расширенный текст планеты в доме или None."""
    refs = _load_extended_refs()
    entry = refs.get('planets_in_houses', {}).get(planet_key, {}).get(str(house_num))
    return _format_ext_entry(entry)


def get_element_text(element_key, mode='dominant'):
    """Возвращает расширенный текст по стихии (mode: dominant/lacking/balanced)."""
    refs = _load_extended_refs()
    e = refs.get('elements_modalities_hemispheres', {}).get('elements', {}).get(element_key)
    if not e:
        return None
    arch = e.get('archetype', '')
    detail = e.get(mode, '')
    return f"{arch} {detail}".strip() or None


def get_modality_text(modality_key, mode='dominant'):
    """Возвращает расширенный текст по кресту."""
    refs = _load_extended_refs()
    m = refs.get('elements_modalities_hemispheres', {}).get('modalities', {}).get(modality_key)
    if not m:
        return None
    arch = m.get('archetype', '')
    detail = m.get(mode, '')
    return f"{arch} {detail}".strip() or None


# ─── СТИЛЬ (палитра МерКаБа + астрологические акценты) ──────────────────────

COLOR_TITLE = RGBColor(0xD4, 0xA0, 0x17)     # золото
COLOR_H1 = RGBColor(0x2C, 0x3E, 0x50)        # тёмно-серый
COLOR_H2 = RGBColor(0x29, 0x80, 0xB9)        # синий-акцент
COLOR_MUTED = RGBColor(0x7F, 0x8C, 0x8D)     # серый для подписей
COLOR_HR = RGBColor(0xBD, 0xC3, 0xC7)        # линия-разделитель


def setup_styles(doc: Document):
    """Calibri 11pt, интерлиньяж 1.15, поля 2см, палитра."""
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    # Title
    title = styles['Title']
    title.font.name = 'Calibri'
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = COLOR_TITLE

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_H1
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_H2
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True


def add_hr(doc):
    """Горизонтальная линия через border нижнего края пустого параграфа."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BDC3C7')
    pBdr.append(bottom)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def set_cell_width(cell, inches):
    """Форсируем ширину ячейки (python-docx баг)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    # удалить предыдущий w:tcW если был
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcPr.append(tcW)


# ─── СЕКЦИИ ──────────────────────────────────────────────────────────────────

def add_header(doc, meta):
    """Титул: имя + дата/время/место/система."""
    title = doc.add_paragraph('Натальная астрологическая карта', style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(meta['name'])
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_H1

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_str = meta['time'] if meta.get('has_time') else '— (время неизвестно)'
    line = f"{meta['date']}  •  {time_str}  •  {meta['city']}"
    run = info.add_run(line)
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_MUTED

    sys_line = doc.add_paragraph()
    sys_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sys_line.add_run(f"Система: {meta['system_ru']}")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_MUTED
    run.italic = True

    add_hr(doc)


def add_how_to_read(doc, text):
    doc.add_heading('Как читать этот отчёт', level=1)
    doc.add_paragraph(text)


def add_quick_summary(doc, chart):
    """3-колоночная таблица: ключевые точки карты."""
    doc.add_heading('Быстрая сводка', level=1)

    planets = chart['planets']
    dist = chart['distributions']

    rows = [
        ('☉ Солнце',      f"{planets['sun']['sign_ru']} {planets['sun']['degrees']}°"),
        ('☽ Луна',        f"{planets['moon']['sign_ru']} {planets['moon']['degrees']}°"),
    ]
    if chart['meta'].get('has_time') and chart.get('ascendant'):
        rows.append(('↑ Асцендент', f"{chart['ascendant']['sign_ru']} {chart['ascendant']['degrees']}°"))
        rows.append(('⊕ MC',         f"{chart['mc']['sign_ru']} {chart['mc']['degrees']}°"))
    rows.append(('🜂 Доминанта стихий', dist['dominant_element_ru']))

    stel = dist.get('stelliums') or []
    if stel:
        rows.append(('★ Стеллиумы', '; '.join(stel)))

    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        set_cell_width(c0, 1.8)
        set_cell_width(c1, 4.2)
        c0.paragraphs[0].add_run(k).bold = True
        c1.paragraphs[0].add_run(v)
        c0.paragraphs[0].paragraph_format.space_after = Pt(2)
        c1.paragraphs[0].paragraph_format.space_after = Pt(2)


def add_wheel_image(doc, png_path):
    """Встраивает PNG колесо карты."""
    if not png_path or not os.path.exists(png_path):
        return
    doc.add_heading('Колесо карты', level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=Inches(6.0))


def add_sun_moon_asc_summary(doc, text):
    doc.add_heading('Солнце · Луна · Асцендент', level=1)
    intro = doc.add_paragraph()
    r = intro.add_run('Три главных архетипа вашей карты — то, как вы проявляетесь, что чувствуете и через какую призму встречаетесь с миром.')
    r.italic = True
    r.font.color.rgb = COLOR_MUTED
    doc.add_paragraph(text)


def add_planet_block(doc, planet_data, interp_text, sub_heading):
    """Блок одной планеты: заголовок, техданные, интерпретация."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{planet_data['symbol']}  {planet_data['name_ru']} — {planet_data['sign_ru']} {planet_data['degrees']}°")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = COLOR_H2

    meta_parts = []
    if planet_data.get('house'):
        meta_parts.append(f"Дом {planet_data['house']}")
    if planet_data.get('retrograde'):
        meta_parts.append('R℞ ретроград')
    meta_parts.append(f"Элемент: {planet_data['element']}")
    meta_parts.append(f"Крест: {planet_data['modality']}")
    meta_p = doc.add_paragraph()
    mr = meta_p.add_run(' · '.join(meta_parts))
    mr.font.size = Pt(9)
    mr.font.color.rgb = COLOR_MUTED
    mr.italic = True

    if interp_text:
        doc.add_paragraph(interp_text)


def add_planets_section(doc, chart, interp):
    doc.add_heading('Планеты', level=1)

    groups = [
        ('Светила и личные планеты',       ['sun', 'moon', 'mercury', 'venus', 'mars']),
        ('Социальные планеты',             ['jupiter', 'saturn']),
        ('Транссатурновые (поколенческие)', ['uranus', 'neptune', 'pluto']),
        ('Кармические точки',              ['true_north_lunar_node', 'chiron', 'mean_lilith']),
    ]

    planet_interps = (interp or {}).get('planets', {})
    planets = chart['planets']

    for group_title, keys in groups:
        doc.add_heading(group_title, level=2)
        for k in keys:
            if k not in planets:
                continue
            # Приоритет: interp.json (от LLM) → расширенный ref (v2.4) → пусто
            text = planet_interps.get(k)
            if not text:
                text = get_planet_in_sign_text(k, planets[k].get('sign_ru', '')) or ''
            add_planet_block(doc, planets[k], text, group_title)


def add_houses_section(doc, chart, interp):
    if not chart['meta'].get('has_time'):
        return
    houses = chart.get('houses') or {}
    if not houses:
        return
    house_interps = (interp or {}).get('houses', {})

    doc.add_heading('Дома', level=1)
    intro = doc.add_paragraph()
    r = intro.add_run('Каждый из 12 домов — это сфера жизни, окрашенная знаком на куспиде и планетами внутри.')
    r.italic = True
    r.font.color.rgb = COLOR_MUTED

    # Таблица куспидов
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, t in enumerate(['Дом', 'Куспид (знак, градус)', 'Планеты в доме']):
        hdr[i].paragraphs[0].add_run(t).bold = True
        set_cell_width(hdr[i], [0.7, 2.3, 3.0][i])

    # Инвертируем planets → houses
    planets_by_house = {}
    for pk, pv in chart['planets'].items():
        h = pv.get('house')
        if h:
            planets_by_house.setdefault(str(h), []).append(f"{pv['symbol']} {pv['name_ru']}")

    for hk in sorted(houses.keys(), key=lambda x: int(x)):
        h = houses[hk]
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(f"{hk}").bold = True
        row[1].paragraphs[0].add_run(f"{h['sign_ru']} {h['degrees']}°")
        row[2].paragraphs[0].add_run(', '.join(planets_by_house.get(hk, [])) or '—')
        for i in range(3):
            set_cell_width(row[i], [0.7, 2.3, 3.0][i])

    # Интерпретации домов (только тех, где есть интерп)
    for hk in sorted(house_interps.keys(), key=lambda x: int(x)):
        txt = house_interps[hk]
        if not txt:
            continue
        doc.add_heading(f"{hk}-й дом — {houses.get(hk, {}).get('sign_ru', '')}", level=2)
        doc.add_paragraph(txt)


def add_aspects_section(doc, chart, interp):
    doc.add_heading('Ключевые аспекты', level=1)
    intro = doc.add_paragraph()
    r = intro.add_run('Аспекты — это угловые связи между планетами. Чем точнее аспект (меньше орб), тем сильнее его влияние.')
    r.italic = True
    r.font.color.rgb = COLOR_MUTED

    # Топ-10 по точности
    aspects = sorted(chart.get('aspects', []), key=lambda a: a.get('orb', 999))[:10]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, t in enumerate(['Планета 1', 'Аспект', 'Планета 2', 'Орб']):
        hdr[i].paragraphs[0].add_run(t).bold = True
        set_cell_width(hdr[i], [1.8, 1.6, 1.8, 0.8][i])

    for a in aspects:
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(a.get('p1_name_ru', a['p1']))
        row[1].paragraphs[0].add_run(a.get('aspect_ru', a['aspect']))
        row[2].paragraphs[0].add_run(a.get('p2_name_ru', a['p2']))
        row[3].paragraphs[0].add_run(f"{a.get('orb', 0)}°")
        for i in range(4):
            set_cell_width(row[i], [1.8, 1.6, 1.8, 0.8][i])

    # Текстовые интерпретации
    aspect_interps = (interp or {}).get('key_aspects', {})
    for akey, txt in aspect_interps.items():
        if not txt:
            continue
        doc.add_heading(akey.replace('_', ' · ').title(), level=2)
        doc.add_paragraph(txt)


def add_distributions_section(doc, chart, interp_text):
    doc.add_heading('Стихии и модальности', level=1)

    dist = chart['distributions']
    elem_ru = {'fire': 'Огонь', 'earth': 'Земля', 'air': 'Воздух', 'water': 'Вода'}
    mod_ru = {'cardinal': 'Кардинальные', 'fixed': 'Фиксированные', 'mutable': 'Мутабельные'}

    # Стихии
    p1 = doc.add_paragraph()
    p1.add_run('Стихии: ').bold = True
    parts = [f"{elem_ru[k]} — {v}" for k, v in dist['elements'].items()]
    p1.add_run('   '.join(parts))

    # Модальности
    p2 = doc.add_paragraph()
    p2.add_run('Кресты: ').bold = True
    parts = [f"{mod_ru[k]} — {v}" for k, v in dist['modalities'].items()]
    p2.add_run('   '.join(parts))

    # Доминирующая
    p3 = doc.add_paragraph()
    r = p3.add_run(f"Доминирующая стихия: {dist['dominant_element_ru']}")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = COLOR_H2

    if interp_text:
        doc.add_paragraph(interp_text)
    else:
        # v2.4: расширенные fallback-тексты по доминирующей стихии и кресту
        elem_text = get_element_text(dist['dominant_element'], 'dominant')
        if elem_text:
            doc.add_paragraph(elem_text)
        mod_text = get_modality_text(dist['dominant_modality'], 'dominant')
        if mod_text:
            doc.add_paragraph(mod_text)


def add_vedic_section(doc, chart, interp):
    vedic = chart.get('vedic')
    if not vedic:
        return
    vi = (interp or {}).get('vedic', {}) or {}

    doc.add_heading('Ведическая оптика (накшатры)', level=1)
    intro = doc.add_paragraph()
    r = intro.add_run('В сидерическом зодиаке (Лахири) Луна и Асцендент попадают в одну из 27 накшатр — лунных стоянок. Это ключ к глубинному архетипу.')
    r.italic = True
    r.font.color.rgb = COLOR_MUTED

    moon_nak = vedic.get('moon_nakshatra')
    if moon_nak:
        doc.add_heading(f"Луна в накшатре {moon_nak.get('name_ru', '?')}", level=2)
        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"Управитель: {moon_nak.get('ruler', '?')} · Божество: {moon_nak.get('deity', '?')}")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = COLOR_MUTED
        meta_run.italic = True
        if vi.get('moon_nakshatra'):
            doc.add_paragraph(vi['moon_nakshatra'])

    asc_nak = vedic.get('ascendant_nakshatra')
    if asc_nak:
        doc.add_heading(f"Асцендент в накшатре {asc_nak.get('name_ru', '?')}", level=2)
        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"Управитель: {asc_nak.get('ruler', '?')} · Божество: {asc_nak.get('deity', '?')}")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = COLOR_MUTED
        meta_run.italic = True
        if vi.get('ascendant_nakshatra'):
            doc.add_paragraph(vi['ascendant_nakshatra'])


def add_stelliums_section(doc, chart, interp_text):
    stel = (chart.get('distributions') or {}).get('stelliums') or []
    if not stel:
        return
    doc.add_heading('Стеллиумы', level=1)
    intro = doc.add_paragraph()
    r = intro.add_run('Стеллиум — три и более планет в одном знаке или доме. Указывает на мощную концентрацию энергии.')
    r.italic = True
    r.font.color.rgb = COLOR_MUTED

    for s in stel:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(s).bold = True

    if interp_text:
        doc.add_paragraph(interp_text)


def add_practical_advice(doc, advice_list):
    if not advice_list:
        return
    doc.add_heading('Как применять карту на практике', level=1)
    for a in advice_list:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(a)


def add_colophon(doc, meta):
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'Астрологические расчёты произведены с использованием Swiss Ephemeris через библиотеку Kerykeion. '
        'Интерпретации сгенерированы ИИ на основе классических архетипов и не являются медицинским, '
        'финансовым или психологическим советом.'
    )
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED
    r.italic = True

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author.add_run('astro-natal-merkaba · автор: Дмитрий · dimkaklasnyi@gmail.com')
    ar.font.size = Pt(8)
    ar.font.color.rgb = COLOR_MUTED


# ─── ОСНОВНОЙ ФЛОУ ──────────────────────────────────────────────────────────

def render(chart_path, interp_path, out_path):
    chart = json.load(open(chart_path, 'r', encoding='utf-8'))
    interp = json.load(open(interp_path, 'r', encoding='utf-8')) if interp_path and os.path.exists(interp_path) else {}

    doc = Document()
    setup_styles(doc)

    # Шапка
    add_header(doc, chart['meta'])

    # Интро
    intro_text = interp.get('intro_how_to_read') or (
        'Натальная карта — это снимок неба на момент вашего рождения. '
        'Она не предсказывает события, а описывает архетипы, через которые проявляется ваша уникальность. '
        'Каждая планета, знак и дом — это язык. Сначала читайте Солнце, Луну и Асцендент — '
        'это ядро личности, эмоциональный мир и способ встречи с реальностью. '
        'Затем углубляйтесь в отдельные планеты и их связи (аспекты).'
    )
    add_how_to_read(doc, intro_text)

    # Быстрая сводка
    add_quick_summary(doc, chart)

    # Колесо карты
    png_path = chart.get('wheel_png')
    if png_path and not os.path.isabs(png_path):
        # разрешаем относительный путь от chart json
        png_path = os.path.join(os.path.dirname(os.path.abspath(chart_path)), png_path)
    add_wheel_image(doc, png_path)

    add_page_break(doc)

    # Солнце · Луна · Асцендент summary
    sum_text = interp.get('sun_moon_asc_summary')
    if sum_text:
        add_sun_moon_asc_summary(doc, sum_text)

    # Планеты
    add_planets_section(doc, chart, interp)

    add_page_break(doc)

    # Дома
    add_houses_section(doc, chart, interp)

    # Аспекты
    add_aspects_section(doc, chart, interp)

    add_page_break(doc)

    # Стихии и модальности
    add_distributions_section(doc, chart, interp.get('elements_modalities'))

    # Стеллиумы
    add_stelliums_section(doc, chart, interp.get('stelliums'))

    # Ведическая оптика
    add_vedic_section(doc, chart, interp)

    # Практические советы
    add_practical_advice(doc, interp.get('practical_advice') or [])

    # Колофон
    add_colophon(doc, chart['meta'])

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description='Рендер DOCX из натальной карты')
    ap.add_argument('--chart', required=True, help='Путь к chart JSON (от build_chart.py)')
    ap.add_argument('--interp', default=None, help='Путь к interpretations JSON (от LLM)')
    ap.add_argument('--out', default=None, help='Путь DOCX (default: <chart>.docx)')
    args = ap.parse_args()

    chart_path = os.path.abspath(args.chart)
    if not os.path.exists(chart_path):
        print(f"❌ chart JSON не найден: {chart_path}")
        sys.exit(1)

    interp_path = os.path.abspath(args.interp) if args.interp else None

    out_path = args.out or chart_path.replace('.json', '.docx')

    try:
        result = render(chart_path, interp_path, out_path)
        print(f"  ✅ DOCX: {result}")
    except Exception as e:
        print(f"❌ Ошибка рендера: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(2)


if __name__ == '__main__':
    main()
