"""
DOCX-рендер отчёта по астероидам (Phase 8 / v1.7.0).
Палитра МерКаБа из render_docx (DRY).
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from render_docx import (  # type: ignore
    setup_styles,
    add_hr,
    add_page_break,
    set_cell_width,
    add_extended_interp_section,
    COLOR_TITLE,
    COLOR_H1,
    COLOR_H2,
    COLOR_MUTED,
)
from render_transits_docx import NATURE_LABEL  # type: ignore


HOUSE_THEMES = {
    1: "Я, тело, самопроявление",
    2: "Деньги, ресурсы, ценности",
    3: "Коммуникация, окружение, обучение",
    4: "Дом, семья, корни",
    5: "Творчество, дети, удовольствия",
    6: "Здоровье, работа, рутина",
    7: "Партнёрство, отношения",
    8: "Трансформация, общие ресурсы",
    9: "Мировоззрение, путешествия, философия",
    10: "Карьера, общественная роль",
    11: "Друзья, цели, будущее",
    12: "Уединение, психея, скрытое",
}


def add_header(doc: Document, meta: dict):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f"Астероиды — {meta['natal_meta']['name']}")
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = COLOR_TITLE
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Церера · Паллада · Юнона · Веста · Хирон")
    sr.font.size = Pt(14)
    sr.font.color.rgb = COLOR_H2
    sub.paragraph_format.space_after = Pt(8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nm = meta["natal_meta"]
    ir = info.add_run(
        f"Натал: {nm['date']} {nm.get('time', '')} · {nm['city']}"
    )
    ir.font.size = Pt(10)
    ir.font.color.rgb = COLOR_MUTED
    ir.italic = True
    info.paragraph_format.space_after = Pt(12)

    add_hr(doc)


def add_intro(doc: Document):
    doc.add_heading("Как читать этот отчёт", level=1)
    p = doc.add_paragraph(
        "Астероиды Церера, Паллада, Юнона и Веста плюс кентавр Хирон — это «фемининная четвёрка» "
        "плюс архетип раненого целителя. Они дополняют картину классических планет, фокусируясь "
        "на конкретных психологических темах: материнство и забота (Церера), стратегический "
        "интеллект (Паллада), партнёрство и обязательства (Юнона), сакральное призвание (Веста), "
        "родовая травма и дар через боль (Хирон). Учитывайте их там, где основная карта оставляет "
        "вопросы или требует уточнения по этим темам."
    )
    p.paragraph_format.space_after = Pt(8)


def add_summary_table(doc: Document, asteroids: dict):
    doc.add_heading("Сводка положений", level=1)

    rows = list(asteroids.values())
    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.4, 1.6, 1.0, 1.0, 1.2]
    headers = ["Астероид", "Знак", "Градус", "Дом", "Движение"]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
        set_cell_width(cell, widths[i])

    for row_i, ast in enumerate(rows, start=1):
        cells = table.rows[row_i].cells
        cells[0].text = f"{ast['symbol']} {ast['name_ru']}"
        cells[1].text = ast["sign_ru"]
        cells[2].text = f"{ast['degrees']}°"
        cells[3].text = f"дом {ast['house']}"
        cells[4].text = "ретроград ℞" if ast["retrograde"] else "директ"
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_asteroid_blocks(doc: Document, asteroids: dict):
    add_page_break(doc)
    doc.add_heading("Архетипы астероид в карте", level=1)

    intro = doc.add_paragraph(
        "Для каждой астероиды ниже — два слоя интерпретации. Первый: окрашивание архетипа "
        "стихией знака (огонь / земля / воздух / вода). Второй: сфера жизни через квадрант "
        "дома (1–3 — самость, 4–6 — близкий круг, 7–9 — другие, 10–12 — коллектив)."
    )
    intro.paragraph_format.space_after = Pt(10)

    for key, ast in asteroids.items():
        # Заголовок астероиды
        h2 = doc.add_paragraph()
        h2.style = doc.styles["Heading 2"]
        h2.paragraph_format.keep_with_next = True
        retro = " ℞" if ast["retrograde"] else ""
        title_run = h2.add_run(
            f"{ast['symbol']} {ast['name_ru']} в {ast['sign_ru']} {ast['degrees']}° "
            f"(дом {ast['house']}{retro})"
        )
        title_run.font.color.rgb = COLOR_H2

        # Архетипный девиз
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.keep_with_next = True
        mr = meta_p.add_run(ast["archetype"])
        mr.font.size = Pt(10)
        mr.italic = True
        mr.font.color.rgb = COLOR_MUTED

        # По элементу
        if ast.get("archetype_in_element"):
            p = doc.add_paragraph()
            p.add_run("Стихия знака. ").bold = True
            p.add_run(ast["archetype_in_element"])
            p.paragraph_format.space_after = Pt(4)

        # По квадранту дома
        if ast.get("archetype_in_quadrant"):
            p = doc.add_paragraph()
            p.add_run("Квадрант дома. ").bold = True
            p.add_run(ast["archetype_in_quadrant"])
            p.paragraph_format.space_after = Pt(4)

        # Тема дома
        house_theme = HOUSE_THEMES.get(ast["house"])
        if house_theme:
            p = doc.add_paragraph()
            p.add_run("Сфера дома. ").bold = True
            p.add_run(f"Дом {ast['house']} — {house_theme}.")
            p.paragraph_format.space_after = Pt(10)


def add_aspects(doc: Document, aspects: list):
    if not aspects:
        return
    add_page_break(doc)
    doc.add_heading("Аспекты к натальным планетам", level=1)
    p = doc.add_paragraph(
        "Орбы для астероид взяты узкие (1–2°), потому что вне точного аспекта астероид редко "
        "проявляется заметно. Перечислены только тесные конфигурации."
    )
    p.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=len(aspects) + 1, cols=4)
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.7, 2.5, 1.4, 0.9]
    headers = ["Астероид", "Натал", "Природа", "Орб"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
        set_cell_width(cell, widths[i])

    for row_i, asp in enumerate(aspects, start=1):
        cells = table.rows[row_i].cells
        cells[0].text = asp["asteroid_name"]
        cells[1].text = f"{asp['aspect_symbol']} {asp['natal_name']}"
        cells[2].text = NATURE_LABEL.get(asp["nature"], "—")
        cells[3].text = f"{asp['orb']}°"
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_practical_advice(doc: Document):
    add_page_break(doc)
    doc.add_heading("Как применять астероиды на практике", level=1)
    advice = [
        "Церера показывает, через что вы заботитесь о близких и о себе. Если она в напряжении — "
        "вероятно, в роду были потери и навык «отпускать» развивался через боль.",
        "Паллада — ваш стиль стратегического мышления. Знак показывает «как» вы анализируете, "
        "дом — «где» применяете эту способность.",
        "Юнона раскрывает запрос к партнёру. Если она конфликтует с натальной Венерой/Марсом — "
        "идеал партнёра отличается от реальной притягивающей энергии. Это нормально, важно знать оба.",
        "Веста показывает призвание и сакральный фокус. Где она стоит — туда стоит вкладываться "
        "глубоко и регулярно, даже если кажется «непрактичным».",
        "Хирон — родовая рана и одновременно главный дар. Тема, в которой вы проходите через боль "
        "и потом помогаете другим. Не лечится, но интегрируется через принятие.",
    ]
    for a in advice:
        bullet = doc.add_paragraph(a, style="List Bullet")
        bullet.paragraph_format.space_after = Pt(4)


def add_colophon(doc: Document):
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Положения астероид рассчитаны через Swiss Ephemeris (kerykeion sweph data, "
        "файл seas_18.se1, период 1800–2400). Орбы аспектов подобраны как у профессиональной "
        "западной школы: 1–2° для астероид, чтобы избежать ложных конфигураций."
    )
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED
    r.italic = True

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run("astro-natal-merkaba · автор: Дмитрий · dimkaklasnyi@gmail.com")
    ar.font.size = Pt(8)
    ar.font.color.rgb = COLOR_MUTED


def render_asteroids_docx(result: dict, out_path: str, interp: dict | None = None) -> str:
    doc = Document()
    setup_styles(doc)

    add_header(doc, result["meta"])
    add_intro(doc)
    add_summary_table(doc, result["asteroids"])
    add_asteroid_blocks(doc, result["asteroids"])
    add_aspects(doc, result["aspects"])
    add_practical_advice(doc)
    if interp:
        add_extended_interp_section(doc, interp)
    add_colophon(doc)

    out_p = Path(out_path)
    out_p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_p))
    return str(out_p)


if __name__ == "__main__":
    import argparse
    import json as _json

    p = argparse.ArgumentParser()
    p.add_argument("--asteroids", required=True)
    p.add_argument("--interp", required=False)
    p.add_argument("--out", required=True)
    args = p.parse_args()

    with open(args.asteroids) as f:
        data = _json.load(f)
    interp = None
    if args.interp and os.path.exists(args.interp):
        with open(args.interp) as f:
            interp = _json.load(f)
    out = render_asteroids_docx(data, args.out, interp=interp)
    print(f"📄 DOCX: {out}")
