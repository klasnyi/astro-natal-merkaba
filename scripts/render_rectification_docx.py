"""DOCX-рендер отчёта по ректификации (Phase 11 / v1.10.0)."""
from __future__ import annotations

import os
import sys
from pathlib import Path

sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from render_docx import (  # type: ignore
    setup_styles, add_hr, add_page_break, set_cell_width,
    COLOR_TITLE, COLOR_H1, COLOR_H2, COLOR_MUTED,
)


def add_header(doc, meta):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nm = meta["natal_meta"]
    tr = title.add_run(f"Ректификация — {nm['name']}")
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = COLOR_TITLE
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Восстановление времени рождения по событиям жизни")
    sr.font.size = Pt(14)
    sr.font.color.rgb = COLOR_H2
    sub.paragraph_format.space_after = Pt(8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    w = meta["window"]
    ir = info.add_run(
        f"Натал: {nm['date']} · {nm['city']}    "
        f"Окно: {w['start']} → {w['end']}, шаг {w['step_minutes']} мин ({meta['candidates_total']} кандидатов)"
    )
    ir.font.size = Pt(10)
    ir.font.color.rgb = COLOR_MUTED
    ir.italic = True
    info.paragraph_format.space_after = Pt(12)

    add_hr(doc)


def add_intro(doc, events):
    doc.add_heading("Как читать этот отчёт", level=1)
    p = doc.add_paragraph(
        "Время рождения, отличающееся всего на 4 минуты, смещает Асцендент примерно на 1°. "
        "Поэтому, если время неизвестно, его можно восстановить — поискать, какое время даёт "
        "карту, в которой натальные точки (особенно ASC, MC и cusps домов) резонируют с "
        "транзитами в дни значимых событий жизни. Этот отчёт перебирает кандидаты времени "
        "и оценивает каждый по сумме «попаданий» транзитов в натальные углы и планеты на даты "
        "событий. Чем точнее транзит на ангулярную точку (1° орбит) — тем выше вес. Топ-3 "
        "кандидата — наиболее вероятные времена."
    )
    p.paragraph_format.space_after = Pt(8)

    doc.add_heading("Использованные события", level=2)
    for e in events:
        bullet = doc.add_paragraph(f"{e['date']} — {e.get('type', '?')}", style="List Bullet")
        bullet.paragraph_format.space_after = Pt(2)


def add_top_candidates_table(doc, top):
    add_page_break(doc)
    doc.add_heading("Топ кандидатов времени", level=1)
    p = doc.add_paragraph(
        "Время с самым высоким score — наиболее вероятное реальное время рождения. "
        "Но всегда сверяйте с памятью свидетелей: ректификация — это статистическая оценка, не доказательство."
    )
    p.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=len(top) + 1, cols=5)
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [0.5, 1.0, 1.0, 1.7, 1.7]
    headers = ["#", "Время", "Score", "ASC", "MC"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
        set_cell_width(cell, widths[i])

    for rank, r in enumerate(top, start=1):
        cells = table.rows[rank].cells
        cells[0].text = str(rank)
        cells[1].text = r["time"]
        cells[2].text = str(r["score"])
        cells[3].text = f"{r['asc']['sign_ru']} {r['asc']['degrees']}°"
        cells[4].text = f"{r['mc']['sign_ru']} {r['mc']['degrees']}°"
        if rank == 1:
            for cell in cells:
                for para in cell.paragraphs:
                    for rn in para.runs:
                        rn.font.bold = True
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_top_candidate_breakdown(doc, top):
    if not top:
        return
    best = top[0]
    add_page_break(doc)
    doc.add_heading(f"Лучший кандидат: {best['time']} (score {best['score']})", level=1)
    p = doc.add_paragraph(
        "Подробная разбивка резонанса по каждому событию. Чем больше совпадений с тесными "
        "орбами — тем сильнее это время «отвечает» на жизненный сценарий."
    )
    p.paragraph_format.space_after = Pt(10)

    for breakdown in best.get("breakdowns", []):
        ev = breakdown["event"]
        h2 = doc.add_paragraph()
        h2.style = doc.styles["Heading 2"]
        h2.paragraph_format.keep_with_next = True
        h2.add_run(f"{ev['date']} — {ev.get('type', '?')}  (score {breakdown['score']})")

        if not breakdown["top_aspects"]:
            p2 = doc.add_paragraph("Нет тесных аспектов в этом окне.")
            p2.paragraph_format.space_after = Pt(6)
            continue

        for asp in breakdown["top_aspects"]:
            line = doc.add_paragraph(
                f"транзит {asp['transit_planet']} → натальный {asp['natal_target']}, "
                f"орб {asp['orb']}°, score {asp['score']}",
                style="List Bullet",
            )
            line.paragraph_format.space_after = Pt(2)


def add_practical_advice(doc):
    add_page_break(doc)
    doc.add_heading("Как использовать результат", level=1)
    advice = [
        "Топ-1 — кандидат с самым высоким резонансом. Проверь его, построив натал в этом времени и сравнив "
        "темы домов с реальной жизнью (где у тебя 4-й дом — что в семье и доме; 10-й — карьера; 7-й — отношения).",
        "Если топ-1 и топ-2 близки по score (разница <10%) — возможно реальное время между ними. Используй "
        "оба для последующих транзитов и прогрессий, выбери то, что лучше работает на практике.",
        "Чем больше событий ты дал на вход (4-7 разнотемных) — тем точнее работает алгоритм. Только свадьба "
        "недостаточно, нужны и потери, и переезды, и карьерные сдвиги.",
        "Уже свадьба, развод, потеря близкого, переезд и карьерное событие — пять разных архетипов планет, "
        "это золотой минимум.",
        "Алгоритм не учитывает прогрессии и солнечные дуги (это тоже хорошие методы) — он опирается только "
        "на транзиты к натальным углам/планетам. Дополни ректификацию вручную: проверь, попадают ли "
        "прогрессии Луны на натальные точки в дни событий.",
        "Время с шагом 8-12 мин — практичный компромисс между точностью и скоростью. Для финальной "
        "проверки можно перезапустить с шагом 4 мин в окне ±15 мин вокруг лучшего кандидата.",
    ]
    for a in advice:
        bullet = doc.add_paragraph(a, style="List Bullet")
        bullet.paragraph_format.space_after = Pt(4)


def add_colophon(doc):
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Ректификация — статистическая оценка через резонанс транзитов с натальными точками "
        "на даты событий. Веса транзитных планет подобраны по тематической близости к типу события. "
        "Это инструмент гипотез, не доказательство. Сверяйте с реальной жизнью."
    )
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED
    r.italic = True

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run("astro-natal-simond · автор: Дмитрий · dimkaklasnyi@gmail.com")
    ar.font.size = Pt(8)
    ar.font.color.rgb = COLOR_MUTED


def render_rectification_docx(result: dict, out_path: str) -> str:
    doc = Document()
    setup_styles(doc)

    add_header(doc, result["meta"])
    add_intro(doc, result["events"])
    add_top_candidates_table(doc, result["top_candidates"])
    add_top_candidate_breakdown(doc, result["top_candidates"])
    add_practical_advice(doc)
    add_colophon(doc)

    out_p = Path(out_path)
    out_p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_p))
    return str(out_p)


if __name__ == "__main__":
    import argparse, json as _json

    p = argparse.ArgumentParser()
    p.add_argument("--rect", required=True)
    p.add_argument("--out", required=True)
    args = p.parse_args()
    with open(args.rect) as f:
        data = _json.load(f)
    out = render_rectification_docx(data, args.out)
    print(f"📄 DOCX: {out}")
