"""
DOCX-рендер отчёта по Vimshottari Dasha (Phase 7 / v1.9.0).
Палитра МерКаБа из render_docx.
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from render_docx import (  # type: ignore
    setup_styles,
    add_hr,
    add_page_break,
    set_cell_width,
    add_extended_interp_section,
    COLOR_TITLE,
    COLOR_H1,
    COLOR_H2,
    COLOR_MUTED,
)


def add_header(doc, meta):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nm = meta["natal_meta"]
    tr = title.add_run(f"Ведические Даши — {nm['name']}")
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = COLOR_TITLE
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Vimshottari Dasha — 120-летний цикл планетных периодов")
    sr.font.size = Pt(14)
    sr.font.color.rgb = COLOR_H2
    sub.paragraph_format.space_after = Pt(8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(
        f"Натал: {nm['date']} {nm.get('time', '')} · {nm['city']}    "
        f"Lahiri ayanamsha: {meta.get('ayanamsha_lahiri', '?')}°    "
        f"На дату: {meta.get('today', '?')}"
    )
    ir.font.size = Pt(10)
    ir.font.color.rgb = COLOR_MUTED
    ir.italic = True
    info.paragraph_format.space_after = Pt(12)

    add_hr(doc)


def add_intro(doc):
    doc.add_heading("Как читать этот отчёт", level=1)
    p = doc.add_paragraph(
        "Vimshottari Dasha — главная предсказательная система ведической астрологии. Жизнь делится "
        "на 9 неравных «эпох» (mahadasha), каждой управляет одна из планет: Кету (7 лет), Венера "
        "(20), Солнце (6), Луна (10), Марс (7), Раху (18), Юпитер (16), Сатурн (19), Меркурий (17). "
        "Полный цикл — 120 лет. Точка старта определяется по натальной Луне в накшатре (одной из 27 "
        "лунных стоянок) — её владыка задаёт первую mahadasha и ту её часть, которая уже прошла "
        "к моменту рождения. Внутри каждой mahadasha идут 9 antardasha (под-периодов) пропорциональной "
        "длины. Архетип владыки — это тема целой эпохи жизни."
    )
    p.paragraph_format.space_after = Pt(8)


def add_nakshatra_block(doc, nakshatra, moon_sid_lon):
    doc.add_heading("Натальная накшатра Луны", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ("Sidereal долгота Луны", f"{moon_sid_lon}°"),
        ("Накшатра", f"{nakshatra['index']}: {nakshatra['name_ru']} ({nakshatra['name_en']})"),
        ("Владыка накшатры", nakshatra["lord_ru"]),
        ("Доля пройдена", f"{nakshatra['fraction_through']*100:.2f}%"),
    ]
    widths = [2.5, 4.0]
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        for r in cells[0].paragraphs[0].runs:
            r.font.bold = True
        for j, w in enumerate(widths):
            set_cell_width(cells[j], w)


def add_mahadasha_table(doc, mahadashas, current_mahadasha):
    add_page_break(doc)
    doc.add_heading("Mahadasha — главные эпохи жизни", level=1)
    p = doc.add_paragraph(
        "Полный 120-летний цикл от момента рождения. Каждая mahadasha — это «эпоха», "
        "тематически окрашенная её владыкой. Текущая выделена."
    )
    p.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=len(mahadashas) + 1, cols=4)
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.5, 1.4, 1.4, 0.8]
    headers = ["Владыка", "Начало", "Конец", "Лет"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
        set_cell_width(cell, widths[i])

    for row_i, m in enumerate(mahadashas, start=1):
        cells = table.rows[row_i].cells
        is_current = current_mahadasha and m["lord"] == current_mahadasha["lord"] and m["start"] == current_mahadasha["start"]
        marker = " ◀" if is_current else ""
        cells[0].text = f"{m['symbol']} {m['lord_ru']}{marker}"
        cells[1].text = m["start"]
        cells[2].text = m["end"]
        cells[3].text = str(m["years"])
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)
        if is_current:
            for cell in cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.font.bold = True


def add_current_focus(doc, current_mahadasha, current_antardasha):
    if not current_mahadasha:
        return
    add_page_break(doc)
    doc.add_heading(f"Сейчас: {current_mahadasha['symbol']} {current_mahadasha['lord_ru']} ({current_mahadasha['start']} → {current_mahadasha['end']})", level=1)

    theme_p = doc.add_paragraph(current_mahadasha["theme"])
    theme_p.paragraph_format.space_after = Pt(10)

    if current_antardasha:
        h2 = doc.add_paragraph()
        h2.style = doc.styles["Heading 2"]
        h2.add_run(
            f"Антардаша: {current_antardasha['symbol']} {current_antardasha['lord_ru']} "
            f"({current_antardasha['start']} → {current_antardasha['end']}, "
            f"{current_antardasha['months']} мес)"
        )

        sub_p = doc.add_paragraph(
            "Антардаша — под-период внутри текущей махадаши. Тема антардаши накладывается на тему "
            "махадаши: например, Юпитер-махадаша × Венера-антардаша = расширение возможностей через "
            "отношения / искусство / эстетику. Чем ближе планеты по природе, тем мягче переход."
        )
        sub_p.paragraph_format.space_after = Pt(8)


def add_antardasha_breakdown(doc, current_mahadasha):
    """Полный список 9 antardasha внутри текущей mahadasha (для планирования)."""
    if not current_mahadasha:
        return
    # Импортируем функцию из build_dashas
    sys.path.insert(0, os.path.dirname(__file__))
    from build_dashas import build_antardashas  # type: ignore

    antardashas = build_antardashas(current_mahadasha)

    doc.add_heading("Все антардаши текущей махадаши", level=1)
    p = doc.add_paragraph(
        "Под-периоды внутри текущей эпохи. Каждый — около нескольких месяцев или лет, "
        "со своим тематическим окрашиванием."
    )
    p.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=len(antardashas) + 1, cols=4)
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.5, 1.4, 1.4, 1.0]
    headers = ["Антардаша", "Начало", "Конец", "Месяцев"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
        set_cell_width(cell, widths[i])

    for row_i, a in enumerate(antardashas, start=1):
        cells = table.rows[row_i].cells
        cells[0].text = f"{a['symbol']} {a['lord_ru']}"
        cells[1].text = a["start"]
        cells[2].text = a["end"]
        cells[3].text = f"{a['months']}"
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_practical_advice(doc):
    add_page_break(doc)
    doc.add_heading("Как читать даши на практике", level=1)
    advice = [
        "Mahadasha — главная тема десятков лет. Если сейчас Юпитер-махадаша — основная сцена жизни "
        "про учительство, рост, мудрость, путешествия и философию. Если Сатурн — про дисциплину, "
        "ответственность, медленное мастерство.",
        "Antardasha — субтема. Внутри Юпитер-махадаши Сатурн-антардаша = «расширение через дисциплину» "
        "(можно учиться чему-то требующему годового труда). Раху-антардаша = «амбиции внутри роста» (легко "
        "переоценить, осторожно с риском).",
        "Резкие переходы между mahadasha (например, Юпитер → Сатурн) часто ощущаются как смена эпохи: "
        "то, что давало рост, теперь требует структуры. Подготовь себя за 1-2 года.",
        "Ведическая система предполагает sidereal зодиак с аянамшей Лахири — он отстаёт от тропического "
        "примерно на 24°. Поэтому накшатры Луны могут не совпадать с тем знаком, к которому привыкли в "
        "западной традиции — это нормально.",
        "Дашa — это потенциал, а не приговор. Тяжёлая Сатурн-mahadasha может быть периодом гения и "
        "мастерства, если работать с её требованиями честно.",
    ]
    for a in advice:
        bullet = doc.add_paragraph(a, style="List Bullet")
        bullet.paragraph_format.space_after = Pt(4)


def add_colophon(doc):
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Расчёт через Swiss Ephemeris (kerykeion sweph + pyswisseph) с аянамшей Lahiri. "
        "Sidereal долгота Луны → накшатра → её владыка → старт цикла Vimshottari. "
        "Mahadasha длительности фиксированы (Vimshottari = 120 лет полный цикл)."
    )
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED
    r.italic = True

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run("astro-natal-simond · автор: Дмитрий · dimkaklasnyi@gmail.com")
    ar.font.size = Pt(8)
    ar.font.color.rgb = COLOR_MUTED


def render_dashas_docx(result: dict, out_path: str, interp: dict | None = None) -> str:
    doc = Document()
    setup_styles(doc)

    add_header(doc, result["meta"])
    add_intro(doc)
    add_nakshatra_block(doc, result["nakshatra"], result["moon_sidereal"]["abs_pos"])
    add_mahadasha_table(doc, result["mahadashas"], result.get("current_mahadasha"))
    add_current_focus(doc, result.get("current_mahadasha"), result.get("current_antardasha"))
    add_antardasha_breakdown(doc, result.get("current_mahadasha"))
    add_practical_advice(doc)
    if interp:
        add_extended_interp_section(doc, interp)
    add_colophon(doc)

    out_p = Path(out_path)
    out_p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_p))
    return str(out_p)


if __name__ == "__main__":
    import argparse
    import json as _json

    p = argparse.ArgumentParser()
    p.add_argument("--dashas", required=True)
    p.add_argument("--interp", required=False)
    p.add_argument("--out", required=True)
    args = p.parse_args()

    with open(args.dashas) as f:
        data = _json.load(f)
    interp = None
    if args.interp and os.path.exists(args.interp):
        with open(args.interp) as f:
            interp = _json.load(f)
    out = render_dashas_docx(data, args.out, interp=interp)
    print(f"📄 DOCX: {out}")
