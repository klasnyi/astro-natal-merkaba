#!/usr/bin/env python3.11
"""
astro-natal-merkaba: render_progressions_docx.py
Рендерит DOCX-отчёт прогрессий (палитра МерКаБа).
"""
import argparse
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("❌ python-docx не установлен", file=sys.stderr)
    sys.exit(1)

from render_docx import (
    setup_styles, add_hr, add_page_break, set_cell_width,
    COLOR_TITLE, COLOR_H1, COLOR_H2, COLOR_MUTED,
)
from render_transits_docx import (
    INTENSITY_COLOR, NATURE_LABEL, MONTH_RU, format_date_ru,
)


def add_header(doc, prog):
    meta = prog['meta']
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f"Прогрессии {meta['natal_name']}")
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = COLOR_TITLE
    title.paragraph_format.space_after = Pt(2)

    age = meta.get('age_years', 0)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    years = int(age)
    months = int((age - years) * 12)
    sr = sub.add_run(f"возраст {years} лет {months} месяцев")
    sr.font.size = Pt(14)
    sr.font.color.rgb = COLOR_H2
    sub.paragraph_format.space_after = Pt(8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(
        f"Натал: {format_date_ru(meta.get('natal_date'))} · "
        f"{meta.get('natal_city', '?')} · "
        f"Прогрессированная дата: {format_date_ru(meta.get('progressed_date'))}"
    )
    ir.font.size = Pt(10)
    ir.font.color.rgb = COLOR_MUTED
    ir.italic = True
    info.paragraph_format.space_after = Pt(12)

    add_hr(doc)


def add_how_to_read(doc, text):
    if not text:
        return
    doc.add_heading('Как читать этот отчёт', level=1)
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_sun_block(doc, prog, sun_text):
    """Прогрессированное Солнце — ключевая секция."""
    doc.add_heading('Прогрессированное Солнце — фокус личности', level=1)
    sun = prog['progressed_planets'].get('sun', {})

    info = doc.add_paragraph()
    info.paragraph_format.keep_with_next = True
    ir = info.add_run(
        f"☉ Прогрессированное Солнце сейчас в {sun.get('sign_ru', '?')} {sun.get('degrees', '?')}°"
    )
    ir.font.size = Pt(12)
    ir.font.bold = True
    info.paragraph_format.space_after = Pt(8)

    if sun_text:
        p = doc.add_paragraph(sun_text)
        p.paragraph_format.space_after = Pt(10)
    else:
        p = doc.add_paragraph(
            f"(Интерпретация прогрессированного Солнца не сгенерирована — "
            f"заполни поле 'progressed_sun' в interp.json.)"
        )
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = COLOR_MUTED


def add_moon_block(doc, prog, moon_text, phase_text):
    """Прогрессированная Луна + лунная фаза."""
    doc.add_heading('Прогрессированная Луна — эмоциональная погода', level=1)
    moon = prog['progressed_planets'].get('moon', {})
    phase = prog.get('progressed_moon_phase', {})

    info = doc.add_paragraph()
    info.paragraph_format.keep_with_next = True
    ir = info.add_run(
        f"☽ Прогрессированная Луна сейчас в {moon.get('sign_ru', '?')} {moon.get('degrees', '?')}°"
    )
    ir.font.size = Pt(12)
    ir.font.bold = True
    info.paragraph_format.space_after = Pt(4)

    # Phase summary
    phase_p = doc.add_paragraph()
    phase_p.paragraph_format.keep_with_next = True
    pr = phase_p.add_run(
        f"Фаза цикла: {phase.get('phase_name', '?')} (разделение {phase.get('separation_degrees', '?')}°). "
        f"{phase.get('phase_meaning', '')}"
    )
    pr.font.size = Pt(10)
    pr.font.color.rgb = COLOR_MUTED
    pr.italic = True
    phase_p.paragraph_format.space_after = Pt(8)

    if moon_text:
        p = doc.add_paragraph(moon_text)
        p.paragraph_format.space_after = Pt(8)
    else:
        p = doc.add_paragraph(
            "(Интерпретация прогрессированной Луны не сгенерирована — "
            "заполни поле 'progressed_moon' в interp.json.)"
        )
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = COLOR_MUTED

    if phase_text:
        h2 = doc.add_paragraph()
        h2.style = doc.styles['Heading 2']
        h2.add_run('О текущей фазе цикла')
        p = doc.add_paragraph(phase_text)
        p.paragraph_format.space_after = Pt(10)


def add_ingresses(doc, prog, ingress_text):
    """Смены знаков прогрессированными планетами."""
    ingresses = prog.get('ingresses', [])
    if not ingresses:
        return
    doc.add_heading('Смены знаков (ingresses)', level=1)
    if ingress_text:
        p = doc.add_paragraph(ingress_text)
        p.paragraph_format.space_after = Pt(8)
    for ing in ingresses:
        line = doc.add_paragraph(style='List Bullet')
        r = line.add_run(
            f"{ing['symbol']} {ing['planet_ru']}: {ing['note']}"
        )
        r.font.size = Pt(11)
        line.paragraph_format.space_after = Pt(4)


def add_inner_planets(doc, prog, inner_text):
    """Прогрессированные внутренние планеты (Mercury, Venus, Mars) — кратко."""
    planets = prog.get('progressed_planets', {})
    inner_keys = ['mercury', 'venus', 'mars']
    has_data = any(planets.get(k) for k in inner_keys)
    if not has_data:
        return

    doc.add_heading('Прогрессированные внутренние планеты', level=1)

    if inner_text:
        p = doc.add_paragraph(inner_text)
        p.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.4, 1.6, 1.4, 1.6]
    headers = ['Планета', 'Прогрессирована в', 'Натальная позиция', 'Сдвиг']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_width(cell, widths[i])
        for r in cell.paragraphs[0].runs:
            r.font.bold = True

    # Натальные позиции достанем из meta для сравнения — но у нас нет в JSON
    # Вместо этого просто выведем прогрессированную позицию + retrograde flag
    for row_i, key in enumerate(inner_keys, start=1):
        p = planets.get(key, {})
        cells = table.rows[row_i].cells
        cells[0].text = f"{p.get('symbol', '')} {p.get('name_ru', key)}"
        cells[1].text = f"{p.get('sign_ru', '?')} {p.get('degrees', '?')}°"
        cells[2].text = '—'  # натальная позиция в JSON прогрессий не передаётся
        cells[3].text = '℞ ретроград' if p.get('retrograde') else 'директ'
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_aspects_to_natal(doc, prog, aspects_interp):
    """Аспекты прогрессий → натал."""
    aspects = prog.get('aspects_to_natal', [])
    if not aspects:
        # Показываем что такое возможно
        doc.add_heading('Аспекты прогрессий к наталу', level=1)
        p = doc.add_paragraph(
            'В этот период прогрессированные планеты не образуют точных аспектов '
            '(в очень узких орбах прогрессий ≤1°) к натальным точкам. Это нормально — '
            'прогрессии работают слоями, аспекты появляются и сходят медленно. '
            'Сейчас доминирующие темы — прогрессированное Солнце и Луна.'
        )
        return

    add_page_break(doc)
    doc.add_heading('Аспекты прогрессий к наталу', level=1)
    p = doc.add_paragraph(
        'Прогрессированные аспекты — это символические ключевые моменты жизни. '
        'Орбы у них узкие (≤1°), поэтому каждый встречающийся аспект — значимое событие.'
    )
    p.paragraph_format.space_after = Pt(8)

    for asp in aspects:
        # H2 с цветом по природе аспекта
        h2 = doc.add_paragraph()
        h2.style = doc.styles['Heading 2']
        h2.paragraph_format.keep_with_next = True
        title_run = h2.add_run(
            f"{asp['progressed_symbol']} прогр. {asp['progressed_planet_ru']} "
            f"{asp['aspect_symbol']} натал. {asp['natal_symbol']} {asp['natal_point_ru']}"
        )
        # Цвет по природе аспекта
        nature = asp.get('aspect_nature', 'neutral')
        if nature == 'tense':
            title_run.font.color.rgb = INTENSITY_COLOR['high']
        elif nature == 'harmonious':
            title_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)  # зелёный
        else:
            title_run.font.color.rgb = COLOR_H2

        # Подзаголовок
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.keep_with_next = True
        natal_house = asp.get('natal_house')
        meta_text = (
            f"Прогр. {asp['progressed_planet_ru']} в {asp['progressed_sign_ru']} "
            f"{asp['progressed_degrees']}° → натал. {asp['natal_point_ru']} "
            f"в {asp['natal_sign_ru']} {asp['natal_degrees']}°"
        )
        if natal_house:
            meta_text += f" (дом {natal_house})"
        meta_text += f" · орб {asp['orb']}° · {NATURE_LABEL.get(nature, 'нейтральный')}"
        if asp.get('progressed_retrograde'):
            meta_text += ' · ℞'
        mr = meta_p.add_run(meta_text)
        mr.font.size = Pt(10)
        mr.font.color.rgb = COLOR_MUTED
        mr.italic = True
        meta_p.paragraph_format.space_after = Pt(4)

        # Текст
        text = aspects_interp.get(asp['key'])
        if text:
            tp = doc.add_paragraph(text)
            tp.paragraph_format.space_after = Pt(10)
        else:
            tp = doc.add_paragraph(
                f"(Интерпретация для {asp['progressed_planet_ru']} {asp['aspect_ru']} "
                f"{asp['natal_point_ru']} не сгенерирована — заполни поле '{asp['key']}' в interp.json.)"
            )
            tp.runs[0].italic = True
            tp.runs[0].font.color.rgb = COLOR_MUTED


def add_practical_advice(doc, advice_list):
    if not advice_list:
        return
    add_page_break(doc)
    doc.add_heading('Что важно в этот символический период', level=1)
    p = doc.add_paragraph(
        'Прогрессии описывают не события, а внутреннее символическое движение. '
        'Эти ориентиры — про то, на что направить внимание сейчас.'
    )
    p.paragraph_format.space_after = Pt(8)
    for advice in advice_list:
        bullet = doc.add_paragraph(advice, style='List Bullet')
        bullet.paragraph_format.space_after = Pt(4)


def add_colophon(doc, prog):
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'Прогрессии рассчитаны через Swiss Ephemeris (kerykeion) методом '
        'дня-за-год. Орбы намеренно узкие (≤1°): прогрессионные аспекты редки и потому значимы. '
        'Интерпретации — архетипические, не предсказания событий.'
    )
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED
    r.italic = True

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run('astro-natal-merkaba · автор: Дмитрий · dimkaklasnyi@gmail.com')
    ar.font.size = Pt(8)
    ar.font.color.rgb = COLOR_MUTED


def render(prog_path, interp_path, out_path):
    with open(prog_path, 'r', encoding='utf-8') as f:
        prog = json.load(f)

    if interp_path and Path(interp_path).exists():
        with open(interp_path, 'r', encoding='utf-8') as f:
            interp = json.load(f)
    else:
        interp = {}

    doc = Document()
    setup_styles(doc)

    add_header(doc, prog)
    add_how_to_read(doc, interp.get('intro_how_to_read'))
    add_sun_block(doc, prog, interp.get('progressed_sun'))
    add_moon_block(doc, prog, interp.get('progressed_moon'), interp.get('moon_phase'))
    add_ingresses(doc, prog, interp.get('ingresses'))
    add_inner_planets(doc, prog, interp.get('inner_planets'))
    add_aspects_to_natal(doc, prog, interp.get('aspects', {}))
    add_practical_advice(doc, interp.get('practical_advice', []))
    add_colophon(doc, prog)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main():
    p = argparse.ArgumentParser()
    p.add_argument('--progressions', required=True)
    p.add_argument('--interp')
    p.add_argument('--out', required=True)
    args = p.parse_args()
    out = render(args.progressions, args.interp, args.out)
    print(f"  📄 DOCX: {out}")


if __name__ == '__main__':
    main()
