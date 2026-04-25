#!/usr/bin/env python3.11
"""
astro-natal-simond: render_composite_docx.py
DOCX-отчёт по композиту (палитра МерКаБа).
"""
import argparse
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("❌ python-docx не установлен", file=sys.stderr)
    sys.exit(1)

from render_docx import (
    setup_styles, add_hr, add_page_break, set_cell_width,
    COLOR_TITLE, COLOR_H1, COLOR_H2, COLOR_MUTED,
)
from render_transits_docx import NATURE_LABEL, format_date_ru
from render_solar_docx import HOUSE_THEMES


def add_header(doc, comp):
    meta = comp['meta']
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f"Композит: {meta['person1_name']} × {meta['person2_name']}")
    tr.font.size = Pt(24)
    tr.font.bold = True
    tr.font.color.rgb = COLOR_TITLE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('Карта отношений как третьей сущности')
    sr.font.size = Pt(13)
    sr.font.color.rgb = COLOR_H2
    sr.italic = True
    sub.paragraph_format.space_after = Pt(8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(
        f"{meta['person1_name']}: {format_date_ru(meta.get('person1_date'))} · "
        f"{meta['person2_name']}: {format_date_ru(meta.get('person2_date'))}\n"
        f"Метод: midpoint composite · Дома: {meta.get('house_system', 'Whole Sign')}"
    )
    ir.font.size = Pt(10)
    ir.font.color.rgb = COLOR_MUTED
    ir.italic = True
    info.paragraph_format.space_after = Pt(12)
    add_hr(doc)


def add_how_to_read(doc, text):
    if not text:
        return
    doc.add_heading('Как читать этот отчёт', level=1)
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_core_summary(doc, comp, summary_text):
    doc.add_heading('Ядро отношений', level=1)
    if summary_text:
        p = doc.add_paragraph(summary_text)
        p.paragraph_format.space_after = Pt(10)

    sun = comp.get('composite_planets', {}).get('sun', {})
    moon = comp.get('composite_planets', {}).get('moon', {})
    asc = comp.get('composite_angles', {}).get('ascendant', {})
    mc = comp.get('composite_angles', {}).get('mc', {})

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ('☉ Солнце композита',
         f"{sun.get('sign_ru', '?')} {sun.get('degrees', '?')}° · "
         f"дом {sun.get('house')} ({HOUSE_THEMES.get(sun.get('house'), '—')})"),
        ('☽ Луна композита',
         f"{moon.get('sign_ru', '?')} {moon.get('degrees', '?')}° · "
         f"дом {moon.get('house')} ({HOUSE_THEMES.get(moon.get('house'), '—')})"),
        ('↑ ASC композита',
         f"{asc.get('sign_ru', '?')} {asc.get('degrees', '?')}°"),
        ('⊕ MC композита',
         f"{mc.get('sign_ru', '?')} {mc.get('degrees', '?')}°"),
    ]
    for i, (lbl, val) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = lbl
        cells[1].text = val
        for r in cells[0].paragraphs[0].runs:
            r.font.bold = True
        set_cell_width(cells[0], 2.3)
        set_cell_width(cells[1], 4.2)


def add_planets_table(doc, comp, planets_text):
    add_page_break(doc)
    doc.add_heading('Все планеты композита', level=1)
    if planets_text:
        p = doc.add_paragraph(planets_text)
        p.paragraph_format.space_after = Pt(10)

    planets = comp.get('composite_planets', {})
    if not planets:
        return
    table = doc.add_table(rows=len(planets) + 1, cols=3)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [2.0, 2.5, 2.0]
    for i, h in enumerate(['Планета', 'Знак · градус', 'Дом']):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_width(cell, widths[i])
        for r in cell.paragraphs[0].runs:
            r.font.bold = True

    for row_i, (k, p) in enumerate(planets.items(), start=1):
        cells = table.rows[row_i].cells
        cells[0].text = f"{p.get('symbol', '')} {p.get('name_ru', k)}"
        cells[1].text = f"{p.get('sign_ru', '?')} {p.get('degrees', '?')}°"
        cells[2].text = f"Дом {p.get('house', '—')} ({HOUSE_THEMES.get(p.get('house'), '—')})"
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_aspects_section(doc, comp, aspects_text):
    aspects = comp.get('aspects', [])
    if not aspects:
        return
    add_page_break(doc)
    doc.add_heading('Аспекты внутри композита', level=1)
    if aspects_text:
        p = doc.add_paragraph(aspects_text)
        p.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=min(len(aspects), 20) + 1, cols=3)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [3.2, 1.5, 1.3]
    for i, h in enumerate(['Аспект', 'Природа', 'Орб']):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_width(cell, widths[i])
        for r in cell.paragraphs[0].runs:
            r.font.bold = True

    for row_i, asp in enumerate(aspects[:20], start=1):
        cells = table.rows[row_i].cells
        cells[0].text = (
            f"{asp['planet1_symbol']} {asp['planet1_ru']} "
            f"{asp['aspect_symbol']} {asp['planet2_symbol']} {asp['planet2_ru']}"
        )
        cells[1].text = NATURE_LABEL.get(asp.get('aspect_nature', 'neutral'), '—')
        cells[2].text = f"{asp['orb']}°"
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_practical_advice(doc, advice_list):
    if not advice_list:
        return
    add_page_break(doc)
    doc.add_heading('Что важно понимать про эти отношения', level=1)
    p = doc.add_paragraph(
        'Композит описывает отношения как живой организм с собственным '
        'характером. Эти ориентиры — про то, что просит развития именно в этой паре.'
    )
    p.paragraph_format.space_after = Pt(8)
    for advice in advice_list:
        bullet = doc.add_paragraph(advice, style='List Bullet')
        bullet.paragraph_format.space_after = Pt(4)


def add_colophon(doc, comp):
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'Композит рассчитан методом midpoint (Robert Hand): каждая позиция = '
        'кратчайшая середина дуги между планетами двух людей. '
        'Дома — Whole Sign от composite ASC. '
        'Это упрощённая, но осмысленная система для midpoint composite.'
    )
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED
    r.italic = True

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run('astro-natal-simond · автор: Дмитрий · dimkaklasnyi@gmail.com')
    ar.font.size = Pt(8)
    ar.font.color.rgb = COLOR_MUTED


def render(comp_path, interp_path, out_path):
    with open(comp_path, 'r', encoding='utf-8') as f:
        comp = json.load(f)
    if interp_path and Path(interp_path).exists():
        with open(interp_path, 'r', encoding='utf-8') as f:
            interp = json.load(f)
    else:
        interp = {}

    doc = Document()
    setup_styles(doc)

    add_header(doc, comp)
    add_how_to_read(doc, interp.get('intro_how_to_read'))
    add_core_summary(doc, comp, interp.get('summary'))
    add_planets_table(doc, comp, interp.get('planets'))
    add_aspects_section(doc, comp, interp.get('aspects'))
    add_practical_advice(doc, interp.get('practical_advice', []))
    add_colophon(doc, comp)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main():
    p = argparse.ArgumentParser()
    p.add_argument('--composite', required=True)
    p.add_argument('--interp')
    p.add_argument('--out', required=True)
    args = p.parse_args()
    out = render(args.composite, args.interp, args.out)
    print(f"  📄 DOCX: {out}")


if __name__ == '__main__':
    main()
