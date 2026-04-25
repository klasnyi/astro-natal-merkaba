"""
DOCX-рендер отчёта по релокации (Phase 9 / v1.8.0).
Палитра МерКаБа из render_docx (DRY).
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from render_docx import (  # type: ignore
    setup_styles,
    add_hr,
    add_page_break,
    set_cell_width,
    COLOR_TITLE,
    COLOR_H1,
    COLOR_H2,
    COLOR_MUTED,
)


HOUSE_THEMES = {
    1: "Я, тело, самопроявление",
    2: "Деньги, ресурсы, ценности",
    3: "Коммуникация, окружение, обучение",
    4: "Дом, семья, корни",
    5: "Творчество, дети, удовольствия",
    6: "Здоровье, работа, рутина",
    7: "Партнёрство, отношения",
    8: "Трансформация, общие ресурсы",
    9: "Мировоззрение, путешествия, философия",
    10: "Карьера, общественная роль",
    11: "Друзья, цели, будущее",
    12: "Уединение, психея, скрытое",
}


def add_header(doc: Document, meta: dict):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nm = meta["natal_meta"]
    tr = title.add_run(f"Релокация {nm['name']} → {meta['target_city']}")
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = COLOR_TITLE
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Жизнь в этом городе через объектив натальной карты")
    sr.font.size = Pt(14)
    sr.font.color.rgb = COLOR_H2
    sub.paragraph_format.space_after = Pt(8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(
        f"Натал: {nm['date']} {nm.get('time', '')} · {nm['city']}    "
        f"Релокация: {meta['target_city']} (lat {meta['target_lat']:.2f}, lon {meta['target_lon']:.2f})"
    )
    ir.font.size = Pt(10)
    ir.font.color.rgb = COLOR_MUTED
    ir.italic = True
    info.paragraph_format.space_after = Pt(12)

    add_hr(doc)


def add_intro(doc: Document):
    doc.add_heading("Как читать этот отчёт", level=1)
    p = doc.add_paragraph(
        "Релокация — это пересчёт натальной карты для другого города при сохранении даты, времени "
        "и UT-момента рождения. Натальные планеты по своим знакам и градусам не меняются — это "
        "ваша внутренняя суть. Но меняется ASC, MC и распределение планет по домам — потому что "
        "дома зависят от долготы места. Релокация показывает, как ваша внутренняя карта проявится "
        "в этом конкретном городе: какие сферы жизни активизируются, какие планеты «выйдут на угол» "
        "и станут особенно ощутимыми, какие темы будут резонировать с этим местом."
    )
    p.paragraph_format.space_after = Pt(8)


def add_angles_comparison(doc: Document, natal: dict, reloc: dict):
    doc.add_heading("Сравнение углов: натал vs релокация", level=1)

    table = doc.add_table(rows=3, cols=3)
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.5, 2.5, 2.5]

    headers = ["Точка", "Натал", "Релокация"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
        set_cell_width(cell, widths[i])

    rows = [
        ("↑ Асцендент",
         f"{natal['ascendant']['sign_ru']} {natal['ascendant']['degrees']}°",
         f"{reloc['ascendant']['sign_ru']} {reloc['ascendant']['degrees']}°"),
        ("⊕ MC",
         f"{natal['mc']['sign_ru']} {natal['mc']['degrees']}°",
         f"{reloc['mc']['sign_ru']} {reloc['mc']['degrees']}°"),
    ]
    for row_i, (label, n_val, r_val) in enumerate(rows, start=1):
        cells = table.rows[row_i].cells
        cells[0].text = label
        cells[1].text = n_val
        cells[2].text = r_val
        for r in cells[0].paragraphs[0].runs:
            r.font.bold = True
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    nr = note.add_run(
        "ASC — «маска» и тело, как вы предъявляете себя миру; меняется в каждом городе. "
        "MC — социальная роль, карьера, призвание; меняется тоже."
    )
    nr.font.size = Pt(10)
    nr.italic = True
    nr.font.color.rgb = COLOR_MUTED


def add_house_changes(doc: Document, changes: list):
    add_page_break(doc)
    doc.add_heading("Планеты, сменившие дом", level=1)
    if not changes:
        p = doc.add_paragraph(
            "В этом городе ни одна планета не сменила натальный дом. Это редкая ситуация — "
            "обычно сдвиг хотя бы у одной-двух планет есть. Релокация ближе к натальной."
        )
        p.paragraph_format.space_after = Pt(8)
        return

    p = doc.add_paragraph(
        "Главный практический вывод релокации — какие сферы жизни активизируются. "
        "Если планета перешла в новый дом, тема этой планеты будет проявляться через темы нового дома. "
        "Особое внимание планетам, которые перешли в угловые дома (1, 4, 7, 10) — они становятся видимыми."
    )
    p.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=len(changes) + 1, cols=4)
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.4, 1.0, 1.0, 3.0]
    headers = ["Планета", "Натал", "Релокация", "Что меняется"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
        set_cell_width(cell, widths[i])

    for row_i, c in enumerate(changes, start=1):
        cells = table.rows[row_i].cells
        cells[0].text = f"{c['symbol']} {c['name_ru']}"
        cells[1].text = f"дом {c['natal_house']}"
        cells[2].text = f"дом {c['reloc_house']}"
        new_theme = HOUSE_THEMES.get(c['reloc_house'], '?')
        cells[3].text = new_theme
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_angular_planets(doc: Document, angular: list):
    if not angular:
        return
    doc.add_heading("Угловые планеты в релокации", level=1)
    p = doc.add_paragraph(
        "Планеты в 1, 4, 7 и 10 домах — самые видимые и активные в этом месте. "
        "Они становятся ключевыми силами, формирующими опыт жизни здесь."
    )
    p.paragraph_format.space_after = Pt(10)

    for ap in angular:
        h2 = doc.add_paragraph()
        h2.style = doc.styles["Heading 2"]
        h2.paragraph_format.keep_with_next = True
        title_run = h2.add_run(
            f"{ap['symbol']} {ap['name_ru']} в доме {ap['house']} "
            f"({HOUSE_THEMES.get(ap['house'], '—')})"
        )
        title_run.font.color.rgb = COLOR_H2

        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.keep_with_next = True
        text = f"{ap['sign_ru']} {ap['degrees']}°"
        if ap.get("retrograde"):
            text += " ℞"
        mr = meta_p.add_run(text)
        mr.font.size = Pt(10)
        mr.italic = True
        mr.font.color.rgb = COLOR_MUTED


def add_house_distribution(doc: Document, planets: dict):
    add_page_break(doc)
    doc.add_heading("Планеты по домам в релокации", level=1)
    by_house = {}
    for k, p_data in planets.items():
        h = p_data.get("house")
        if h is None:
            continue
        by_house.setdefault(h, []).append(p_data)

    for house_num in range(1, 13):
        plist = by_house.get(house_num, [])
        if not plist:
            continue
        h2 = doc.add_paragraph()
        h2.style = doc.styles["Heading 2"]
        h2.paragraph_format.keep_with_next = True
        h2.add_run(f"Дом {house_num} — {HOUSE_THEMES.get(house_num, '?')}")

        line = ", ".join(
            f"{p['symbol']} {p['name_ru']} ({p['sign_ru']} {p['degrees']}°"
            + (" ℞" if p.get("retrograde") else "") + ")"
            for p in plist
        )
        para = doc.add_paragraph(line)
        para.paragraph_format.space_after = Pt(6)


def add_practical_advice(doc: Document):
    add_page_break(doc)
    doc.add_heading("Что важно знать про этот город", level=1)
    advice = [
        "Релокация показывает энергетический ландшафт жизни в конкретном месте. Это не «лучше / хуже» "
        "натальной карты — это другой акцент.",
        "Угловые планеты в релокации становятся центральными силами. Если в натале они в кадентных домах, "
        "релокация может «выпустить» их потенциал на видимый план.",
        "Если важная планета (Sun, Moon, Saturn, ASC-управитель) перешла в 12 дом релокации — место может "
        "ощущаться как «уединённое», подходящее для внутренней работы, не для социального проявления.",
        "Если Венера или Юпитер перешли в угловые дома — место «удачное» для отношений или роста.",
        "Релокация не отменяет натальную карту — она даёт второй слой. Аспекты между планетами (доминирующие "
        "конфигурации) остаются прежними.",
        "Решение о переезде — больше чем астрология. Используйте релокацию как один из факторов наряду с "
        "карьерой, отношениями, финансами и интуицией.",
    ]
    for a in advice:
        bullet = doc.add_paragraph(a, style="List Bullet")
        bullet.paragraph_format.space_after = Pt(4)


def add_colophon(doc: Document):
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Релокация рассчитана через Swiss Ephemeris (kerykeion). UT-момент рождения сохраняется, "
        "координаты места заменяются на целевой город. ASC, MC и cusps домов пересчитаны для новой долготы. "
        "Натальные планеты по знакам не меняются — меняется только их «сцена» (дома)."
    )
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED
    r.italic = True

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run("astro-natal-merkaba · автор: Дмитрий · dimkaklasnyi@gmail.com")
    ar.font.size = Pt(8)
    ar.font.color.rgb = COLOR_MUTED


def render_relocation_docx(result: dict, out_path: str) -> str:
    doc = Document()
    setup_styles(doc)

    add_header(doc, result["meta"])
    add_intro(doc)
    add_angles_comparison(doc, result["natal"], result["relocation"])
    add_house_changes(doc, result["house_changes"])
    add_angular_planets(doc, result["angular_planets"])
    add_house_distribution(doc, result["relocation"]["planets"])
    add_practical_advice(doc)
    add_colophon(doc)

    out_p = Path(out_path)
    out_p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_p))
    return str(out_p)


if __name__ == "__main__":
    import argparse
    import json as _json

    p = argparse.ArgumentParser()
    p.add_argument("--reloc", required=True)
    p.add_argument("--out", required=True)
    args = p.parse_args()

    with open(args.reloc) as f:
        data = _json.load(f)
    out = render_relocation_docx(data, args.out)
    print(f"📄 DOCX: {out}")
