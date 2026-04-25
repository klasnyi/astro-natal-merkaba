#!/usr/bin/env python3.11
"""
astro-natal-simond: render_transits_docx.py
Рендерит DOCX-отчёт по транзитам (палитра МерКаБа).

Вход:
  --transits  <path>  transits.json от build_transits.py
  --interp    <path>  интерпретации от Claude (JSON)
  --out       <path>  путь к DOCX

Структура interp JSON:
{
  "intro_how_to_read": "...",
  "summary": "...",
  "active": {"saturn_square_sun": "...", ...},
  "upcoming": "...",
  "stations": "...",
  "practical_advice": ["...", ...]
}
"""
import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("❌ python-docx не установлен", file=sys.stderr)
    sys.exit(1)

import os

# Переиспользуем стили из render_docx (палитра МерКаБа уже настроена там)
from render_docx import (
    setup_styles, add_hr, add_page_break, set_cell_width,
    COLOR_TITLE, COLOR_H1, COLOR_H2, COLOR_MUTED, COLOR_HR,
)


# ─── ЦВЕТА ИНТЕНСИВНОСТИ ────────────────────────────────────────────────

INTENSITY_COLOR = {
    'high': RGBColor(0xC0, 0x39, 0x2B),    # красный — высокая
    'medium': RGBColor(0xE6, 0x7E, 0x22),  # оранжевый — средняя
    'low': RGBColor(0x7F, 0x8C, 0x8D),     # серый — низкая
}

INTENSITY_LABEL = {
    'high': '🔴 Высокая',
    'medium': '🟠 Средняя',
    'low': '⚪ Низкая',
}

NATURE_LABEL = {
    'harmonious': 'гармоничный',
    'tense': 'напряжённый',
    'neutral': 'нейтральный',
}

MOVEMENT_LABEL = {
    'applying': 'нарастает (applying)',
    'separating': 'расходится (separating)',
    'unknown': '—',
}

MONTH_RU = {
    1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
    5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
    9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря',
}


def format_date_ru(date_iso):
    """2026-04-25 → '25 апреля 2026'"""
    if not date_iso:
        return '—'
    try:
        dt = datetime.fromisoformat(date_iso.split('T')[0])
        return f"{dt.day} {MONTH_RU[dt.month]} {dt.year}"
    except (ValueError, KeyError):
        return date_iso


# ─── СЕКЦИИ ─────────────────────────────────────────────────────────────

def add_header(doc, transits):
    """Title + meta (натал и транзит)."""
    meta = transits['meta']
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f"Транзиты {meta['natal_name']}")
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = COLOR_TITLE
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"на {format_date_ru(meta['transit_date'])}")
    sr.font.size = Pt(14)
    sr.font.color.rgb = COLOR_H2
    sub.paragraph_format.space_after = Pt(8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(
        f"Натал: {format_date_ru(meta.get('natal_date'))} · "
        f"{meta.get('natal_city', '?')} · {meta.get('natal_system', 'western')}"
    )
    ir.font.size = Pt(10)
    ir.font.color.rgb = COLOR_MUTED
    ir.italic = True
    info.paragraph_format.space_after = Pt(12)

    add_hr(doc)


def add_how_to_read(doc, text):
    if not text:
        return
    h = doc.add_heading('Как читать этот отчёт', level=1)
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def add_summary_section(doc, transits, summary_text):
    """Quick summary с топ-3 транзитами + общим текстом."""
    doc.add_heading('Текущий период', level=1)

    if summary_text:
        p = doc.add_paragraph(summary_text)
        p.paragraph_format.space_after = Pt(10)

    # Топ-3 таблица
    top = transits.get('active_transits_top', [])[:3]
    if not top:
        doc.add_paragraph('Сейчас нет активных транзитов с орбом ≤3°.')
        return

    doc.add_heading('🔑 Самые сильные транзиты сейчас', level=2)
    table = doc.add_table(rows=len(top) + 1, cols=4)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Транзит', 'Орб', 'Движение', 'Интенсивность']
    widths = [3.0, 0.7, 1.7, 1.4]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_width(cell, widths[i])
        for r in cell.paragraphs[0].runs:
            r.font.bold = True

    for row_i, asp in enumerate(top, start=1):
        cells = table.rows[row_i].cells
        cells[0].text = (
            f"{asp['transit_symbol']} {asp['transit_planet_ru']} "
            f"{asp['aspect_symbol']} {asp['natal_symbol']} {asp['natal_point_ru']}"
        )
        cells[1].text = f"{asp['orb']}°"
        cells[2].text = MOVEMENT_LABEL.get(asp.get('movement', 'unknown'), '—')
        cells[3].text = INTENSITY_LABEL.get(asp.get('intensity', 'low'), '—')
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_biwheel_image(doc, biwheel_path):
    """Встраивает bi-wheel PNG (если есть) после блока 'Текущий период'."""
    if not biwheel_path or not os.path.exists(biwheel_path):
        return
    doc.add_heading('Натал × Транзиты — bi-wheel', level=2)
    intro = doc.add_paragraph()
    r = intro.add_run(
        'Двойной круг: натальная карта в центре, транзиты по внешнему кольцу. '
        'Линии между ними — активные аспекты. Цвета линий: красный (квадрат), '
        'оранжевый (оппозиция), зелёный (трин), синий (секстиль), белый (соединение).'
    )
    r.italic = True
    r.font.color.rgb = COLOR_MUTED
    r.font.size = Pt(9)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(biwheel_path, width=Inches(6.2))
    p.paragraph_format.space_after = Pt(8)


def add_active_transits(doc, transits, active_interp):
    """Каждый активный транзит со своей секцией."""
    add_page_break(doc)
    doc.add_heading('Активные транзиты — детально', level=1)

    p = doc.add_paragraph(
        'Каждый блок ниже — отдельный транзит, активный сейчас. '
        'Они разные по силе и времени действия: транзиты медленных планет '
        '(Сатурн, Уран, Нептун, Плутон) длятся месяцы и работают глубоко; '
        'быстрых (Луна, Меркурий, Венера) — дни и часы.'
    )
    p.paragraph_format.space_after = Pt(12)

    top = transits.get('active_transits_top', [])
    if not top:
        doc.add_paragraph('Активных транзитов с орбом ≤3° сейчас нет.')
        return

    for asp in top:
        # H2 — название транзита с цветом интенсивности
        h2 = doc.add_paragraph()
        h2.style = doc.styles['Heading 2']
        h2.paragraph_format.keep_with_next = True
        title_run = h2.add_run(
            f"{asp['transit_symbol']} {asp['transit_planet_ru']} "
            f"{asp['aspect_symbol']} {asp['natal_symbol']} {asp['natal_point_ru']}"
        )
        title_run.font.color.rgb = INTENSITY_COLOR.get(
            asp.get('intensity', 'low'), COLOR_H2
        )

        # Подзаголовок: позиция + орб + движение + интенсивность
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.keep_with_next = True
        meta_run = meta_p.add_run(
            f"Транзитный {asp['transit_planet_ru']} в {asp['transit_sign_ru']} "
            f"{asp['transit_degrees']}°"
        )
        if asp.get('transit_retrograde'):
            meta_run.add_text(' ℞')
        natal_house = asp.get('natal_house')
        if natal_house:
            meta_run.add_text(f" → натальный {asp['natal_point_ru']} в {asp['natal_sign_ru']} "
                              f"{asp['natal_degrees']}° (дом {natal_house})")
        else:
            meta_run.add_text(f" → натальный {asp['natal_point_ru']} в {asp['natal_sign_ru']} "
                              f"{asp['natal_degrees']}°")
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = COLOR_MUTED
        meta_run.italic = True

        # Краткие параметры
        params_p = doc.add_paragraph()
        params_p.paragraph_format.keep_with_next = True
        params_p.paragraph_format.space_after = Pt(4)
        params_run = params_p.add_run(
            f"Орб: {asp['orb']}° · "
            f"Движение: {MOVEMENT_LABEL.get(asp.get('movement', 'unknown'))} · "
            f"Природа: {NATURE_LABEL.get(asp.get('aspect_nature', 'neutral'))}"
        )
        if asp.get('estimated_exact_date'):
            params_run.add_text(f" · Точность ≈ {format_date_ru(asp['estimated_exact_date'])}")
        params_run.font.size = Pt(9)
        params_run.font.color.rgb = COLOR_MUTED

        # Текст интерпретации
        text = active_interp.get(asp['key'])
        if text:
            tp = doc.add_paragraph(text)
            tp.paragraph_format.space_after = Pt(10)
        else:
            tp = doc.add_paragraph(
                f"Интерпретация для {asp['transit_planet_ru']} "
                f"{asp['aspect_ru']} {asp['natal_point_ru']} не сгенерирована "
                f"(заполни поле '{asp['key']}' в interp.json)."
            )
            tp.runs[0].italic = True
            tp.runs[0].font.color.rgb = COLOR_MUTED
            tp.paragraph_format.space_after = Pt(10)


def add_upcoming(doc, transits, upcoming_text):
    """Таблица ближайших точностей."""
    upcoming = transits.get('upcoming_exactness', [])
    if not upcoming:
        return

    add_page_break(doc)
    doc.add_heading('Ближайшие точности', level=1)

    if upcoming_text:
        p = doc.add_paragraph(upcoming_text)
        p.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=len(upcoming) + 1, cols=3)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [2.4, 3.5, 0.9]
    headers = ['Дата (приблизительно)', 'Транзит', 'Орб сейчас']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_width(cell, widths[i])
        for r in cell.paragraphs[0].runs:
            r.font.bold = True

    for row_i, asp in enumerate(upcoming, start=1):
        cells = table.rows[row_i].cells
        cells[0].text = format_date_ru(asp.get('estimated_exact_date', ''))
        cells[1].text = (
            f"{asp['transit_symbol']} {asp['transit_planet_ru']} "
            f"{asp['aspect_symbol']} {asp['natal_symbol']} {asp['natal_point_ru']}"
        )
        cells[2].text = f"{asp['orb']}°"
        for i, w in enumerate(widths):
            set_cell_width(cells[i], w)


def add_stations(doc, transits, stations_text):
    """Ретроградные планеты сейчас."""
    retro = transits.get('retrograde_now', [])
    if not retro and not stations_text:
        return
    doc.add_heading('Ретроградные планеты сейчас', level=1)

    if retro:
        line = ', '.join(
            f"{r['symbol']} {r['name_ru']} в {r['sign_ru']} {r['degrees']}°"
            for r in retro
        )
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(8)
    else:
        doc.add_paragraph('Ретроградных планет сейчас нет.')

    if stations_text:
        p = doc.add_paragraph(stations_text)
        p.paragraph_format.space_after = Pt(10)


def add_practical_advice(doc, advice_list):
    if not advice_list:
        return
    add_page_break(doc)
    doc.add_heading('Что делать в этот период', level=1)
    p = doc.add_paragraph(
        'Эти советы — не предписания, а ориентиры. Архетипы, '
        'активирующиеся через текущие транзиты, лучше всего работать через сознательное действие.'
    )
    p.paragraph_format.space_after = Pt(8)
    for advice in advice_list:
        bullet = doc.add_paragraph(advice, style='List Bullet')
        bullet.paragraph_format.space_after = Pt(4)


def add_colophon(doc, transits):
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'Транзиты рассчитаны через Swiss Ephemeris (kerykeion). '
        'Оценочные даты точности — линейная экстраполяция; ретроградные петли '
        'могут сдвигать реальную дату на ±несколько дней. '
        'Интерпретации — архетипические, не предсказания событий.'
    )
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_MUTED
    r.italic = True

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run('astro-natal-simond · автор: Дмитрий · dimkaklasnyi@gmail.com')
    ar.font.size = Pt(8)
    ar.font.color.rgb = COLOR_MUTED


# ─── ОСНОВНОЙ ФЛОУ ──────────────────────────────────────────────────────

def render(transits_path, interp_path, out_path):
    with open(transits_path, 'r', encoding='utf-8') as f:
        transits = json.load(f)

    if interp_path and Path(interp_path).exists():
        with open(interp_path, 'r', encoding='utf-8') as f:
            interp = json.load(f)
    else:
        interp = {}

    doc = Document()
    setup_styles(doc)

    # Структура
    add_header(doc, transits)
    add_how_to_read(doc, interp.get('intro_how_to_read'))
    add_summary_section(doc, transits, interp.get('summary'))
    # Bi-wheel PNG (если build_transits.py вызывался с --biwheel)
    biwheel_path = transits.get('biwheel_png')
    if biwheel_path and not os.path.isabs(biwheel_path):
        biwheel_path = os.path.join(
            os.path.dirname(os.path.abspath(transits_path)), biwheel_path
        )
    add_biwheel_image(doc, biwheel_path)
    add_active_transits(doc, transits, interp.get('active', {}))
    add_upcoming(doc, transits, interp.get('upcoming'))
    add_stations(doc, transits, interp.get('stations'))
    add_practical_advice(doc, interp.get('practical_advice', []))
    add_colophon(doc, transits)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main():
    p = argparse.ArgumentParser(description='Рендер DOCX из транзитов')
    p.add_argument('--transits', required=True, help='transits.json')
    p.add_argument('--interp', help='interp.json (опционально — без него заглушки)')
    p.add_argument('--out', required=True, help='Путь к DOCX')
    args = p.parse_args()

    out = render(args.transits, args.interp, args.out)
    print(f"  📄 DOCX: {out}")


if __name__ == '__main__':
    main()
